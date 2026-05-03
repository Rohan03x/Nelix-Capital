"""
gen_doc3.py  —  NIKE Valuation Model: Complete Analysis
Generates "NIKE Valuation - Complete Analysis.docx"
Covers every cell, formula, quirk, and assumption identified across three
analysis passes of the 'valuatione' worksheet + all CIQ raw data tables.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─────────────────────────────────────────────────────────────
#  HELPERS
# ─────────────────────────────────────────────────────────────

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.style.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
    return h

def add_para(doc, text, bold=False, italic=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def add_callout(doc, text, label="NOTE"):
    """Highlighted paragraph that stands out."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(f"[{label}]  {text}")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x0B)  # dark red
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light List Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


# ─────────────────────────────────────────────────────────────
#  DOCUMENT
# ─────────────────────────────────────────────────────────────

doc = Document()

# Title
title = doc.add_heading("NIKE, Inc. — Equity Valuation Model: Complete Analysis", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run(
    "Exhaustive cell-by-cell documentation of the 'valuatione' worksheet\n"
    "NYSE: NKE  |  FY2015–FY2024 historical  |  FY2025–FY2031 forecast\n"
    "Model file: NIKE Valuation strategy.xlsx  —  DO NOT MODIFY"
).bold = True

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS  (narrative)
# ═══════════════════════════════════════════════════════════════
add_heading(doc, "How to Use This Document", 1)
add_para(doc,
    "This document is the authoritative reference for every formula, assumption, "
    "data source, and quirk in the NIKE valuation model. It is organised into nine "
    "parts that follow the logical flow of the model from structure to outputs. "
    "Sections marked [BUG] or [QUIRK] highlight formula inconsistencies or unusual "
    "modelling choices that analysts should be aware of when interrogating results.")
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  PART I — MODEL STRUCTURE
# ═══════════════════════════════════════════════════════════════
add_heading(doc, "PART I — Model Structure and Architecture", 1)

add_heading(doc, "1.1  Worksheet Layout", 2)
add_para(doc,
    "The workbook contains three sheets: 'instructions' (a read-me), 'valuatione' "
    "(the full model, ~590 rows), and 'Sheet1' (CIQ comparable company data). "
    "All calculations live in 'valuatione'. The sheet is not password-protected "
    "but should be treated as read-only to preserve model integrity.")

add_para(doc,
    "Columns are mapped as follows:\n"
    "  • Column A  — row labels / CIQ field names\n"
    "  • Columns B–K  — historical years FY2015 through FY2024 (hard data from CIQ)\n"
    "  • Columns L–R  — forecast years FY2025 through FY2031 (formula-driven)\n"
    "  • Column S/T  — occasional annotation comments or sanity-check values")

add_heading(doc, "1.2  Three-Statement Integration", 2)
add_para(doc,
    "The model is a fully integrated three-statement model. Balance sheet, income "
    "statement, and cash flow statement are linked via the following identities, "
    "each explicitly checked by formula:")
items = [
    ("Cash from Ops (row 445)", "= row 233 (model's CFO)", "L445: =L233"),
    ("Cash from Investing (row 451)", "= row 234 (model's CFI)", "L451: =L234"),
    ("Cash from Financing (row 466)", "= row 235 (model's CFF)", "L466: =L235"),
    ("Net Change in Cash (row 469)", "= CFO+CFI+CFF+FX (rows 445+451+466+468)", "L469: =L445+L451+L466+L468"),
    ("Ending Cash (row 500)", "= prior cash + ΔCash (row 236)", "L500: =K500+L236"),
    ("Total Assets (row 518)", "= Total Liabilities + Equity (row 541)", "L541: =L540+L533"),
    ("Retained Earnings (row 537)", "= prior RE + NI + dividends", "L537: =L361+L464+K537"),
]
add_table(doc,
    ["Statement Link", "Relationship", "Example Formula"],
    items,
    [2.2, 2.5, 2.0])

add_heading(doc, "1.3  Time Axis", 2)
add_para(doc,
    "Row 10 contains year labels (2015 through 2031). The model uses CIQ's "
    "fiscal-year convention: NIKE's fiscal year ends 31 May. So 'FY2024' = "
    "the year ended 31 May 2024. Row 28 stores the column index to identify "
    "the last historical year (K = column 11, FY2024) and the first forecast "
    "year (L = column 12, FY2025). Forecast columns run L through R = 7 years.")

add_heading(doc, "1.4  Data Sources", 2)
add_para(doc,
    "Historical financial data is sourced from S&P Capital IQ (CIQ) in three "
    "raw-data blocks embedded directly in the worksheet:")
add_table(doc,
    ["Block", "Rows", "Content"],
    [
        ("Income Statement", "315–412", "Revenue through SBC detail; 10 fiscal years FY2015–FY2024"),
        ("Cash Flow Statement", "415–478", "Operating/Investing/Financing activities + supplemental items"),
        ("Balance Sheet", "480–590", "Full BS + NOL/FIN48/Fair Value supplemental tables"),
    ],
    [1.5, 1.2, 4.0])
add_para(doc,
    "Each raw-data block carries a header (rows 415–425 for CF, 480–490 for BS) "
    "with metadata: source = 'S&P Capital IQ - Standard', currency = USD, "
    "magnitude = Thousands (K), and period = Custom (FY ended May 31). "
    "The 'CIQ Restatement Type Code' rows (e.g. row 477) show 'NC' (no change) "
    "for most years; 'RS' (restated) appears for the CF statement in FY2016–FY2017 "
    "and for the BS in FY2015–FY2017, meaning those historical figures reflect "
    "subsequent restatements and may differ from originally-filed 10-Ks. "
    "'RUP' (rolled-up) appears in the BS calculation code for FY2015–FY2021.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  PART II — REVENUE AND INCOME STATEMENT
# ═══════════════════════════════════════════════════════════════
add_heading(doc, "PART II — Revenue and Income Statement", 1)

add_heading(doc, "2.1  Revenue Drivers Table (Rows 320–327)", 2)
add_para(doc,
    "Rows 320–327 contain the key assumption inputs that drive all forecast "
    "income statement lines. These are the PRIMARY input cells of the model. "
    "No other assumptions override these except items explicitly hard-coded elsewhere.")
add_table(doc,
    ["Row", "Label", "FY2025 Value", "Description"],
    [
        ("320", "Revenue Growth Rate", "Formula", "= (L338/K338) - 1; this row derives growth from the revenue row itself, so the real driver is row 338"),
        ("321", "COGS % of Revenue", "~57%", "Historical average COGS/Revenue applied to forecast years"),
        ("322", "SG&A % of Revenue", "~31%", "Historical SG&A ratio scaled forward"),
        ("323", "D/TA Ratio (Debt/Total Assets)", "~24%", "Used to compute Long-term Debt in row 528; see Section 7.2"),
        ("324", "Payout Ratio", "-0.4", "Negative sign convention; used as coefficient in buyback formula (row 460). ~40% of NI returned via buybacks"),
        ("325", "CapEx % of Revenue", "~3.2%", "Drives capex in forecast; row 248 = -revenue × capex%"),
        ("326", "D&A % of Revenue", "~2.0%", "Drives depreciation in forecast; row 249"),
        ("327", "NWC % of Revenue", "~8%", "Net Operating Working Capital as % of revenue; drives ΔNOWC"),
    ],
    [0.5, 2.0, 1.3, 3.0])

add_callout(doc,
    "Row 324 stores the payout ratio as a NEGATIVE number (e.g. -0.4 = 40%). "
    "This is because the buyback formula reads: L460 = L361 × L324 where L361 "
    "is Net Income (positive). The product yields a negative buyback cash flow, "
    "consistent with the sign convention that cash outflows are negative. "
    "This means the cell value itself looks like a negative payout ratio, "
    "which can be confusing when reading the drivers table in isolation.", "QUIRK")

add_heading(doc, "2.2  Revenue Formula Chain (Row 338)", 2)
add_para(doc,
    "Historical revenue (columns B–K) is hard data from CIQ. Forecast revenue is:")
add_para(doc, "    L338: =K338 × (1 + L320)^1", bold=True)
add_para(doc,
    "The explicit '^1' exponent is redundant (raises to the power of 1) but "
    "intentional — the analyst may have inserted it for clarity or as a template "
    "for multi-year compounding formulas. Each subsequent forecast year references "
    "the prior year: M338 = L338 × (1 + M320)^1, etc. FY2024 revenue = $51,362m.")
add_table(doc,
    ["Year", "Revenue ($m)", "Growth %"],
    [
        ("FY2015", "27,790", "—"),
        ("FY2016", "32,376", "+16.5%"),
        ("FY2017", "34,350", "+6.1%"),
        ("FY2018", "36,397", "+6.0%"),
        ("FY2019", "39,117", "+7.5%"),
        ("FY2020", "37,403", "-4.4%  (COVID)"),
        ("FY2021", "44,538", "+19.1%  (recovery)"),
        ("FY2022", "46,710", "+4.9%"),
        ("FY2023", "51,217", "+9.6%"),
        ("FY2024", "51,362", "+0.3%  (slowdown)"),
    ],
    [1.0, 1.5, 2.2])

add_heading(doc, "2.3  COGS and Gross Margin (Row 339)", 2)
add_para(doc,
    "L339: =L338 × L321 (revenue × COGS%). FY2024 COGS = $28,925m, gross margin "
    "= 43.6%. The model holds COGS% constant at approximately the FY2024 level "
    "for all forecast years; no margin expansion or compression is modelled. "
    "This is a conservative assumption given NIKE's stated margin recovery targets.")

add_heading(doc, "2.4  SG&A and Operating Expense (Row 342)", 2)
add_para(doc,
    "SG&A is modelled differently from COGS. The forecast formula is:")
add_para(doc, "    L342: =(L342/L338) × M338  [i.e. prior year SG&A% × current year revenue]", bold=True)
add_para(doc,
    "More precisely: M342 = L342/L338 × M338, meaning each year's SG&A is derived "
    "by scaling the PRIOR year's SG&A ratio onto the current year's revenue. "
    "This produces the same numerical result as applying a constant SG&A% but "
    "the formula is chained year-by-year rather than anchored to FY2024. "
    "FY2024 SG&A = $15,829m (30.8% of revenue).")

add_heading(doc, "2.4.1  Selling/Marketing Expense (Row 392)", 3)
add_para(doc,
    "Row 392 breaks out NIKE's selling/marketing expense separately from the "
    "total SG&A reported to CIQ. FY2024 selling expense = $4,285m (8.3% of revenue). "
    "This figure represents demand-creation spending (advertising, sponsorships, "
    "athletes). The model does NOT separately forecast marketing spend — it is "
    "embedded in total SG&A. Analysts should be aware that NIKE's marketing "
    "commitment is contractually large; it cannot be cut proportionally with revenue.")
add_para(doc,
    "Row 393: Net Rental Expense = $1,051m FY2024. This is the pre-ASC 842 "
    "operating lease cash cost, useful for pre/post-IFRS16 comparisons and "
    "for constructing EBITDAR (see Section 2.7).")

add_heading(doc, "2.5  Restructuring Charges (Row 347)", 2)
add_para(doc,
    "Historical restructuring: FY2023 = $402m, FY2024 = $443m. These relate to "
    "NIKE's announced workforce reduction program (~2% headcount, ~1,600 roles) "
    "and product SKU rationalisation. Row 347 formula in forecast: =0 for all "
    "years (restructuring is assumed fully completed by end of FY2024). "
    "This means the model assumes clean, non-recurring results from FY2025 onwards. "
    "Any continuation of restructuring would make actual earnings worse than modelled.")
add_callout(doc,
    "The restructuring freeze has a MATERIAL impact on the DCF. Stripping $443m "
    "of annual charges adds ~$350m after-tax to forecast NOPAT, inflating the "
    "terminal value by roughly $3.5–4bn (at a 10% WACC). If restructuring "
    "represents an ongoing cost of doing business, this is an optimistic assumption.", "IMPORTANT")

add_heading(doc, "2.6  EBIT and Operating Income (Rows 350–355)", 2)
add_para(doc,
    "EBIT = Gross Profit − SG&A − Restructuring (row 350). FY2024 EBIT = $6,754m "
    "(13.1% margin). Below EBIT:\n"
    "  • Row 356: Interest expense on IBD (average IBD × average interest rate; see Section 7.1)\n"
    "  • Row 355: EBT = EBIT + Interest income − Interest expense\n"
    "  • Row 357: =B355/B359 = EBT÷Tax provision = 1÷effective rate (unlabelled sanity check)\n"
    "  • Row 358: =B355×$R$25 = tax at statutory rate for comparison to actual\n"
    "Rows 357 and 358 are hidden/unlabelled check rows that exist for all historical "
    "years but have no forecast equivalents and are invisible in the formatted model.")

add_heading(doc, "2.7  Tax Analysis (Rows 357–389)", 2)
add_para(doc,
    "The model contains an unusually detailed tax section:")
add_table(doc,
    ["Row", "Label", "FY2024 Value", "Notes"],
    [
        ("359", "Tax Provision", "$1,000m", "Accrual-basis; matches reported income tax expense"),
        ("357", "EBT/Tax (hidden check)", "~5.7x", "= 1/effective rate ≈ 1/17.5%"),
        ("358", "Tax at statutory rate", "~$1,417m", "= EBT × 21% (TCJA federal rate)"),
        ("382", "Effective Tax Rate %", "17.5%", "Separate % field; see full history below"),
        ("383", "Current domestic tax", "$983m", "US federal + state current portion"),
        ("384", "Current foreign tax", "$514m", "Non-US current tax"),
        ("385", "Total current tax", "$1,497m", "= row 383 + row 384"),
        ("386", "Deferred domestic tax", "-$483m", "US deferred portion (benefit)"),
        ("387", "Deferred foreign tax", "-$14m", "Non-US deferred"),
        ("388", "Total deferred tax", "-$497m", "= row 386 + row 387"),
        ("359", "Net provision", "$1,000m", "= current + deferred = $1,497m + (-$497m)"),
        ("472", "Cash taxes paid", "$1,299m", "CF statement supplemental; exceeds provision by $299m"),
    ],
    [0.5, 2.0, 1.3, 3.0])
add_para(doc,
    "Cash taxes paid ($1,299m) exceed the provision ($1,000m) because NIKE pays "
    "quarterly estimated taxes that may overshoot the final provision. The difference "
    "creates a prepaid tax asset on the balance sheet.")

add_heading(doc, "2.7.1  Effective Tax Rate History", 3)
add_table(doc,
    ["Year", "ETR%", "Driver"],
    [
        ("FY2015", "27.0%", "Pre-TCJA; normal rate"),
        ("FY2016", "19.3%", "SBC windfalls begin (ASU 2016-09)"),
        ("FY2017", "13.2%", "TCJA transition-year benefits"),
        ("FY2018", "54.6%", "TCJA one-time repatriation toll charge ($2bn+); anomaly year"),
        ("FY2019", "13.4%", "Post-TCJA settled; low effective rate"),
        ("FY2020", "18.2%", "COVID year; lower pre-tax income"),
        ("FY2021", "18.3%", "Normal"),
        ("FY2022", "9.1%", "LOWEST — massive SBC deductions from elevated stock price"),
        ("FY2023", "18.8%", "Normal"),
        ("FY2024", "17.5%", "Slight improvement"),
    ],
    [0.8, 0.8, 5.0])
add_para(doc,
    "FY2022's 9.1% ETR is the outlier: NIKE's stock price was near all-time highs "
    "in mid-2022, meaning RSU vests and stock option exercises produced large tax "
    "deductions (excess tax benefits from SBC = windfall deductions). The model "
    "forecasts a normalised ~18% rate.")
add_para(doc,
    "Row 25 stores the forecast tax rate used in CFADS and other formulas. "
    "The rate in column R25 ($R$25) is the statutory anchor that rows 358, 260, "
    "and 261 attempt to reference — see the CFADS bug in Section 5.2.")

add_heading(doc, "2.8  Net Income and EPS (Rows 361–373)", 2)
add_para(doc,
    "Row 361: Net Income = EBT − Tax. FY2024 = $5,700m. "
    "Row 362: Diluted shares = ~1,503m (FY2024).")
add_table(doc,
    ["Row", "Metric", "FY2024", "Notes"],
    [
        ("363", "Reported diluted EPS", "$3.73", "= NI / diluted shares"),
        ("372", "Normalised diluted EPS", "$2.918", "Strips restructuring, applies normalised tax rate"),
        ("373", "Normalised NI", "$4,464m", "vs $5,700m reported; diff = restructuring + tax normalisation"),
        ("389", "Normalised NI (alt calc)", "$4,464,375k", "Same concept, shown at thousands level"),
    ],
    [0.5, 2.0, 1.3, 3.0])
add_callout(doc,
    "Normalised EPS ($2.92) is 22% BELOW reported EPS ($3.73). This means the "
    "reported P/E significantly understates the 'true' earnings multiple. At a "
    "stock price of $75 (illustrative), reported P/E = 20x but normalised P/E = 26x. "
    "The model uses reported NI for the DCF, which implicitly assumes restructuring "
    "charges do NOT recur — creating an inconsistency (normalised EPS is lower, "
    "but the DCF is built off unstripped NI).", "IMPORTANT")

add_heading(doc, "2.9  EBITDA Variants (Rows 378–381)", 2)
add_table(doc,
    ["Row", "Metric", "FY2024 ($m)", "Formula"],
    [
        ("378", "EBITDA", "7,550", "EBIT + D&A ($796m)"),
        ("379", "EBITA", "6,754", "= EBIT; no goodwill amortisation at NIKE"),
        ("380", "EBIT", "6,754", "Operating income before restructuring add-back"),
        ("381", "EBITDAR", "8,601", "= EBITDA + net rental expense ($1,051m); pre-lease-standard view"),
    ],
    [0.5, 2.2, 1.5, 2.5])
add_para(doc,
    "EBITDAR is the most relevant metric for retail/consumer companies with "
    "significant store footprints. NIKE's $1,051m net rental expense (row 393) "
    "is the pre-ASC 842 equivalent of the on-balance-sheet lease payments now "
    "reported separately. EBITDAR allows comparison to pre-ASC 842 peers.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  PART III — CASH FLOW STATEMENT
# ═══════════════════════════════════════════════════════════════
add_heading(doc, "PART III — Cash Flow Statement", 1)

add_heading(doc, "3.1  Operating Cash Flow Build (Rows 217–233)", 2)
add_para(doc,
    "The model builds its own operating cash flow from first principles rather "
    "than pulling it directly from CIQ. The CIQ reported CFO (row 445) merely "
    "checks back against the model's computed row 233 via: L445 = L233.")
add_table(doc,
    ["Row", "Component", "FY2024 ($m)", "Formula Pattern"],
    [
        ("434", "Net Income (CF basis)", "5,700", "Hard data from CIQ; same as IS"),
        ("437", "D&A Total", "796", "= Property D&A + goodwill/intangible amort"),
        ("438", "Asset writedown (restructuring)", "48", "FY2024 only; FY2023 = $156m; non-cash portion of restructuring"),
        ("439", "Stock-based compensation", "804", "Added back (non-cash)"),
        ("440", "Other operating activities", "-635", "Includes working capital reclasses and non-operating items"),
        ("441", "ΔAccounts Receivable", "-329", "Increase in AR = use of cash"),
        ("442", "ΔInventories", "+908", "Inventory drawdown = source of cash (FY2024 destocking)"),
        ("443", "ΔAccounts Payable", "+397", "Increase in AP = source of cash"),
        ("444", "ΔOther net operating assets", "-260", "Prepaid, accruals etc"),
        ("445", "Total Cash from Operations", "7,429", "Sum; links to model row 233"),
    ],
    [0.5, 2.0, 1.3, 3.0])
add_para(doc,
    "FY2021 CFO = $6,657m (peak) driven by massive NI ($5,727m) and reduced capex "
    "during COVID-recovery. FY2020 CFO = $2,485m (trough) from low NI ($2,539m) "
    "and inventory build ($-1,854m in FY2020).")
add_para(doc,
    "Asset writedown (row 438): FY2023 = $156m and FY2024 = $48m. These are the "
    "non-cash component of NIKE's restructuring — the portion that writes down "
    "inventory and fixed assets rather than paying cash severance. They are added "
    "back in CFO because the expense was already deducted in NI.")

add_heading(doc, "3.2  Investing Activities (Rows 447–451)", 2)
add_table(doc,
    ["Row", "Item", "FY2024 ($m)", "Notes"],
    [
        ("447", "Capital expenditure", "-812", "Gross CapEx; see Section 7.3 for BS-derivation"),
        ("448", "Sale of PP&E", "nil (FY2017–FY2024)", "Only FY2015–FY2017 had small asset sales ($3–13m)"),
        ("449", "Net investment in marketable securities", "+1,721", "Large positive = matured investments >new purchases; FY2021 = -$3,276m (COVID cash deployment)"),
        ("450", "Other investing", "-15", "Acquisitions of small businesses/IP"),
        ("451", "Total Cash from Investing", "+894", "Links to model row 234"),
    ],
    [0.5, 2.0, 1.0, 3.2])
add_para(doc,
    "FY2021's -$3,276m in marketable securities investment is the single largest "
    "year for securities purchases — NIKE deployed its COVID cash hoard into "
    "short-term treasuries and money market instruments, explaining the jump in "
    "short-term investments from $197m (FY2019) to $3,587m (FY2021) on the BS.")

add_heading(doc, "3.3  Financing Activities (Rows 453–466)", 2)
add_para(doc,
    "The financing section has a notable structural feature: historical dividends "
    "are split across two rows for different time periods.")
add_table(doc,
    ["Row", "Item", "Period Coverage", "FY2024 ($m)"],
    [
        ("453–455", "Debt issued (short + long)", "FY2015–FY2024", "nil (no new debt in FY2024)"),
        ("456–458", "Debt repaid (short + long)", "FY2015–FY2024", "nil (no repayment in FY2024; J458=-$504m in FY2023)"),
        ("459", "Issuance of common stock (SBC proceeds)", "FY2015–FY2024", "$667m"),
        ("460", "Repurchase of common stock", "FY2015–FY2031", "-$4,250m FY2024; forecast = NI × payout_ratio"),
        ("461", "Common dividends paid", "FY2015–FY2019 ONLY", "Historical only; FY2019 = -$1,332m"),
        ("462", "Equity raisings (separate row)", "FY2015–FY2031 = 0", "Hard-coded zero throughout"),
        ("463", "Common and/or Pref. dividends", "FY2020–FY2024", "-$2,169m FY2024"),
        ("464", "Total dividends paid", "FY2015–FY2031", "-$2,169m FY2024; forecast = row 460/2"),
        ("465", "Other financing", "FY2015–FY2024", "-$136m FY2024 (debt issuance costs, etc.)"),
        ("466", "Total Cash from Financing", "FY2015–FY2031", "-$5,888m FY2024; links to model row 235"),
    ],
    [0.5, 2.5, 1.8, 1.5])

add_heading(doc, "3.3.1  Dividend Row Split Explained", 3)
add_para(doc,
    "NIKE changed its dividend reporting between FY2019 and FY2020: prior to FY2020 "
    "common dividends were reported on a separate line (row 461); from FY2020 onwards "
    "CIQ consolidates common and any preferred dividends into a single line (row 463). "
    "NIKE has no preferred equity, so rows 461 and 463 are economically identical — "
    "but the two-row structure means row 464 (Total Dividends) is required to sum "
    "both rows to get a continuous series across all years.")

add_heading(doc, "3.3.2  Forecast Dividend Formula (Row 464)", 3)
add_para(doc, "    L464: =L460/2  →  M464: =L464  →  N464: =M464  ...  R464: =Q464", bold=True)
add_para(doc,
    "Two separate mechanics:\n"
    "  1. FY2025 dividends = FY2025 buybacks ÷ 2. Since buybacks = NI × |payout_ratio|, "
    "dividends = NI × |payout_ratio| / 2. If NI = $5,000m and payout = 40%, "
    "buybacks = $2,000m and dividends = $1,000m (total return = $3,000m = 60% of NI).\n"
    "  2. FY2026–FY2031: dividends FROZEN at the FY2025 level (M464 = L464 etc). "
    "Dividends do not grow with NI in the forecast — an unrealistic but "
    "conservative simplifying assumption.")
add_callout(doc,
    "The dividend/buyback ratio of 0.5 (dividends = half of buybacks) is an "
    "approximation of NIKE's historical pattern (~40–50% of total returns in dividends "
    "historically). However, freezing dividends in years 2–7 means the model "
    "implicitly increases the buyback:dividend ratio as NI grows, which may "
    "overstate buybacks relative to dividends in later forecast years.", "QUIRK")

add_heading(doc, "3.3.3  Equity Raisings (Row 462)", 3)
add_para(doc,
    "Row 462 'equity raisings' is hard-coded to 0 for ALL years (both historical and "
    "forecast). This represents new equity issuances beyond SBC proceeds. NIKE has "
    "not issued net new equity since the 1990s; the row is a placeholder that "
    "analysts can activate if modelling a capital raise scenario. "
    "Note: Row 459 (stock issuances = SBC proceeds) is NOT zero — it is the "
    "proceeds from employees exercising options/vesting RSUs, which flows through "
    "APIC and offsets the dilution impact.")

add_heading(doc, "3.4  Supplemental CF Items (Rows 470–476)", 2)
add_table(doc,
    ["Row", "Item", "FY2024 ($m)", "Significance"],
    [
        ("471", "Cash interest paid", "381", "Actual cash outflow for interest; used in CFADS calculation"),
        ("472", "Cash taxes paid", "1,299", "Exceeds provision ($1,000m) by $299m; timing difference"),
        ("473", "CIQ Levered FCF", "5,853", "CIQ formula: CFO − CapEx + asset sales"),
        ("474", "CIQ Unlevered FCF", "6,021", "CIQ formula: ULFCF = LFCF + after-tax interest"),
        ("475", "CIQ Net Working Capital change", "-1,012", "= sum of rows 441–444; consistent with model"),
        ("476", "CIQ Net Debt Issued", "NA (FY2024)", "No net debt change in FY2024; FY2023 = -$504m repayment"),
    ],
    [0.5, 2.2, 1.3, 2.8])
add_para(doc,
    "The model's own FCF metrics differ slightly from CIQ's because the model "
    "uses its own intermediate calculations (EFCF, CFADS) rather than the "
    "CIQ supplemental rows. Section 5 covers these in detail.")

add_heading(doc, "3.5  FX Adjustment (Row 468)", 2)
add_para(doc,
    "Row 468 'Foreign Exchange Rate Adj.' captures the effect of currency translation "
    "on cash balances held in non-USD currencies. FY2024 = -$16m (immaterial). "
    "In the forecast, all L468–R468 are hard-coded to 0, meaning the model assumes "
    "no FX translation effects — appropriate for a base-case USD model.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  PART IV — BALANCE SHEET
# ═══════════════════════════════════════════════════════════════
add_heading(doc, "PART IV — Balance Sheet", 1)

add_heading(doc, "4.1  Current Assets", 2)
add_para(doc,
    "The forecast methodology for each current asset line uses revenue-scaling: "
    "forecast item = (item / FY2024 revenue) × forecast revenue. This assumes "
    "each balance sheet line is a fixed % of revenue, anchored to the FY2024 ratio.")
add_table(doc,
    ["Row", "Item", "FY2024 ($m)", "Forecast Method", "Notes"],
    [
        ("500", "Cash & equivalents", "9,860", "Prior cash + ΔCash (row 236)", "Residual; driven by 3-statement"),
        ("501", "Short-term investments", "1,722", "Revenue-scaled", "Very volatile historically ($197m–$4,423m)"),
        ("502", "Trading securities", "nil", "Hard-coded 0 in forecast", "Only existed FY2015–FY2016 ($78m/$7m)"),
        ("503", "Total cash & ST investments", "11,582", "= sum 500+501+502", ""),
        ("504", "Accounts receivable", "4,427", "Revenue-scaled", "Days sales = ~31 days FY2024"),
        ("505", "Notes receivable", "nil", "Hard-coded 0", "Only FY2015 ($150m)"),
        ("507", "Inventory", "7,519", "Revenue-scaled", "Destocked from $8,454m peak in FY2023"),
        ("508", "Prepaid expenses", "1,224", "Revenue-scaled", "Includes prepaid marketing/sponsorship"),
        ("509", "Other current assets", "630", "FROZEN at FY2024 value", "=K509 in all forecast years"),
        ("510", "Total current assets", "25,382", "Sum of above", ""),
    ],
    [0.5, 1.8, 1.2, 1.8, 1.5])
add_callout(doc,
    "Other Current Assets (row 509) is deliberately frozen rather than revenue-"
    "scaled. This suggests the analyst judged this line to be driven by factors "
    "other than revenue (e.g. short-term derivatives, tax receivables). "
    "Similarly, Short-term borrowings (row 522) and Deferred tax assets LT "
    "(row 516) are frozen.", "QUIRK")

add_heading(doc, "4.2  Non-Current Assets", 2)
add_heading(doc, "4.2.1  Net PP&E — Hard-Coded After FY2025 (Row 513)", 3)
add_para(doc,
    "This is one of the most important and unusual features of the balance sheet model:")
add_para(doc,
    "  • L513 (FY2025): =K513/K$338×L$338  [revenue-scaled normally]\n"
    "  • M513–R513: Hard-coded to $7,885,489,990  (≈$7,885m) in ALL years\n"
    "  • Cell S513 contains the comment: 'Hard coded'\n"
    "  • This means Net PP&E is frozen at ~$7.885bn from FY2026 through FY2031",
    bold=False)
add_callout(doc,
    "The Gross PP&E and Accumulated Depreciation rows (511, 512) are still "
    "revenue-scaled in the forecast, creating an internal inconsistency: "
    "Gross PP&E and AccumDeprec keep growing, but Net PP&E is fixed. "
    "The hard-coded Net PP&E overrides the component-level formulas. "
    "Economically this means CapEx ≈ D&A in perpetuity (maintenance capex only), "
    "which is a reasonable steady-state assumption but is ASSUMED rather than derived.", "IMPORTANT")

add_heading(doc, "4.2.2  PP&E Component Breakdown (Rows 554–558)", 3)
add_table(doc,
    ["Component", "FY2015 ($m)", "FY2024 ($m)", "Change"],
    [
        ("Land", "273", "329", "+20% over 9 years"),
        ("Buildings", "1,250", "3,439", "+175%; major expansion of owned facilities"),
        ("Machinery / Equipment", "3,329", "3,123", "-6%; mature asset base"),
        ("Construction in Progress (CIP)", "350", "193", "-45%; projects completed; $525m in FY2023 came online"),
        ("Leasehold Improvements", "1,150", "2,023", "+76%; store/office buildouts"),
        ("Total Gross PP&E", "6,352", "13,632", "+115%; total doubling"),
    ],
    [2.2, 1.2, 1.2, 2.2])
add_para(doc,
    "CIP dropped from $525m (FY2023) to $193m (FY2024), indicating major capital "
    "projects came online in FY2024 — likely the new technology and distribution "
    "infrastructure NIKE has been building. Buildings more than doubled, reflecting "
    "NIKE's investment in owned supply-chain assets (distribution centres) "
    "post-COVID to reduce dependence on third-party logistics.")

add_heading(doc, "4.2.3  Goodwill (Row 514)", 3)
add_para(doc,
    "Goodwill = $240m FY2024, down from $284m FY2022 peak. Forecast: frozen at "
    "K514 = $240m throughout (L514: =K514, M514: =L514, etc.). No acquisitions "
    "assumed. NIKE's goodwill is small relative to its size, reflecting a history "
    "of brand-building rather than acquisitive growth.")

add_heading(doc, "4.2.4  Other Long-term Assets (Row 517)", 3)
add_para(doc,
    "Row 517 = $2,046m FY2024 (up from $525m FY2015). This is revenue-scaled in "
    "the forecast. It captures right-of-use assets from operating leases (post-ASC 842), "
    "long-term derivatives, and equity investments. The tripling since FY2015 is "
    "almost entirely explained by ASC 842 adoption in FY2020 when NIKE's operating "
    "lease ROU asset ($2,913m) was added to the balance sheet.")

add_heading(doc, "4.2.5  Deferred Tax Assets Long-term (Row 516)", 3)
add_para(doc,
    "Row 516 DTA = $2,465m FY2024. Forecast: FROZEN at FY2024 value (=K516 for "
    "all periods). The analyst treats deferred taxes as static — a reasonable "
    "simplification since DTA/DTL movements are hard to forecast and often "
    "offsetting. DTA jumped from $663m (FY2019) to $1,963m (FY2020) due to "
    "accelerated depreciation deductions and COVID-era loss carrybacks.")

add_heading(doc, "4.3  Current Liabilities", 2)
add_table(doc,
    ["Row", "Item", "FY2024 ($m)", "Forecast Method"],
    [
        ("520", "Accounts payable", "2,851", "Revenue-scaled"),
        ("521", "Accrued expenses", "3,880", "Revenue-scaled"),
        ("522", "Short-term borrowings", "6", "FROZEN at $6m (=K522)"),
        ("523", "Current portion of long-term debt", "1,000", "FROZEN at FY2024 value ($1,000m = 2025 maturity)"),
        ("524", "Current portion of leases", "477", "Revenue-scaled"),
        ("525", "Income taxes payable", "534", "Revenue-scaled"),
        ("526", "Other current liabilities", "1,845", "Revenue-scaled"),
        ("527", "Total current liabilities", "10,593", "Sum"),
    ],
    [0.5, 2.0, 1.2, 2.5])
add_callout(doc,
    "The current portion of LT debt (row 523) is frozen at $1,000m for all "
    "forecast years. In reality this represents the FY2024 maturity of NIKE's "
    "2023 bond. After FY2025 payment, this should theoretically reflect the "
    "next tranche of maturities, not a static $1bn. This is a known simplification.", "QUIRK")

add_heading(doc, "4.4  Long-term Debt — Derivation (Row 528)", 2)
add_para(doc,
    "Long-term debt in the forecast is derived from the D/TA ratio driver (row 323):")
add_para(doc, "    L528: =K518 × K323  (FY2025 debt = FY2024 total assets × FY2024 D/TA ratio)", bold=True)
add_para(doc, "    M528: =L518 × L323  (FY2026 debt = FY2025 total assets × FY2025 D/TA ratio)", bold=True)
add_callout(doc,
    "[BUG] L528 uses K518 (FY2024 total assets) rather than L518 (FY2025 total assets). "
    "This means FY2025 long-term debt is based on PRIOR year assets × CURRENT D/TA ratio, "
    "while all subsequent years (M528 through R528) use current-year assets. "
    "This creates a one-year lag for FY2025 only. The numerical impact is small "
    "if total assets don't change dramatically between FY2024 and FY2025.", "BUG")
add_para(doc,
    "Long-term debt history: $1,079m (FY2015) → $9,406m (FY2020) → $7,934m (FY2024). "
    "NIKE issued $6,134m of LT debt in FY2020 to build a COVID liquidity buffer. "
    "The company has been systematically reducing debt since FY2021. "
    "S528 is annotated 'IBD' marking this as an interest-bearing debt row.")

add_heading(doc, "4.4.1  IBD Classification", 3)
add_para(doc,
    "The following rows are annotated 'IBD' (Interest-Bearing Debt) in column S:")
add_table(doc,
    ["Row", "Component", "FY2024 ($m)"],
    [
        ("522", "Short-term borrowings", "6"),
        ("523", "Current portion of LT debt", "1,000"),
        ("524", "Current portion of leases (IFRS/ASC842)", "477"),
        ("528", "Long-term debt", "7,934"),
        ("529", "Long-term leases", "2,566"),
    ],
    [0.5, 2.5, 1.5])
add_para(doc,
    "Total IBD = $6m + $1,000m + $477m + $7,934m + $2,566m = $11,983m FY2024. "
    "The WACC calculation uses average IBD × interest rate to derive interest expense. "
    "Including lease obligations in IBD means NIKE's leverage looks considerably "
    "higher than a pre-ASC 842 analysis would suggest.")

add_heading(doc, "4.5  Long-term Leases (Row 529)", 2)
add_para(doc,
    "Row 529 = $2,566m FY2024. Revenue-scaled in forecast. Pre-ASC 842 (pre-FY2020) "
    "NIKE had no on-balance-sheet long-term lease liability. The $2,913m that appeared "
    "in FY2020 represented the present value of all future operating lease payments "
    "discounted at NIKE's incremental borrowing rate (~3%). "
    "Note: CIQ row 551 'Debt Equivalent of Operating Leases' shows $8,408m for "
    "FY2024 — the FULL undiscounted obligation. The $5.8bn gap between on-BS ($2,566m) "
    "and CIQ's PV estimate reflects different discount rate assumptions and lease term scope.")

add_heading(doc, "4.6  Other Non-current Liabilities (Rows 530–532)", 2)
add_table(doc,
    ["Row", "Item", "FY2024 ($m)", "Forecast Method"],
    [
        ("530", "Pension & post-retirement benefits", "nil", "HARD-CODED ZERO (=0) for all forecast years; T530 = equity/(IBD+equity) sanity check"),
        ("531", "Deferred tax liability non-current", "145", "Revenue-scaled"),
        ("532", "Other non-current liabilities", "2,442", "Revenue-scaled"),
    ],
    [0.5, 2.2, 1.2, 2.5])
add_para(doc,
    "Cell T530 contains the formula: =K540/(K201+K540)×100. This computes "
    "FY2024 equity as a percentage of (IBD + equity) = the equity weight in "
    "capital structure. The result feeds into the WACC verification. "
    "K201 = total IBD in FY2024; K540 = total equity = $14,430m.")

add_heading(doc, "4.7  Equity Section (Rows 534–540)", 2)
add_para(doc,
    "The equity section has four components, each forecast separately:")
add_table(doc,
    ["Row", "Component", "FY2024 ($m)", "Forecast Formula", "Mechanics"],
    [
        ("535", "Common stock (par value)", "3", "Frozen: =K535", "NIKE's $0.0001 par; essentially nil"),
        ("536", "Additional Paid-In Capital", "13,409", "L536: =K536+L460+L459", "APIC += buybacks (neg) + issuances (pos)"),
        ("537", "Retained Earnings", "965", "L537: =L361+L464+K537", "RE += NI + dividends (neg) + prior RE"),
        ("538", "Accumulated OCI", "53", "Frozen: =K538", "No FX/hedge movements modelled"),
        ("539", "Total Common Equity", "14,430", "= sum 535:538", ""),
        ("540", "Total Equity", "14,430", "=L539", "No minority interest"),
    ],
    [0.5, 2.0, 1.2, 2.2, 2.0])

add_heading(doc, "4.7.1  APIC Mechanics (Row 536)", 3)
add_para(doc,
    "APIC increases each year by: (a) stock issuance proceeds (row 459, always positive) "
    "and (b) share repurchases (row 460, always negative for a buyer). "
    "NIKE does NOT carry treasury stock; repurchased shares are immediately retired, "
    "so all buyback effects hit APIC (reduce it) rather than a separate treasury account. "
    "APIC has grown from $6,773m (FY2015) to $13,409m (FY2024) despite $29bn+ of "
    "buybacks, because SBC issuances and exercises continuously add to APIC.")

add_heading(doc, "4.7.2  Retained Earnings Roll-Forward (Row 537)", 3)
add_para(doc, "    L537: =L361 + L464 + K537", bold=True)
add_para(doc,
    "  • L361 = forecast Net Income (NI adds to RE)\n"
    "  • L464 = forecast dividends (negative, reduces RE)\n"
    "  • K537 = prior year ending RE\n"
    "NIKE's RE has been suppressed by its aggressive buyback program. RE = $965m "
    "in FY2024 despite cumulative NI of ~$40bn+ since FY2015 — nearly all profits "
    "have been returned to shareholders or reinvested in assets.\n"
    "FY2020 RE went NEGATIVE (-$191m), a rare event for an investment-grade company "
    "— the result of COVID-year low NI plus continuation of buybacks and dividends "
    "funded from borrowings.")

add_heading(doc, "4.7.3  Accumulated OCI (Row 538)", 3)
add_para(doc,
    "OCI = $53m frozen in forecast (=K538 for all years). Historically volatile: "
    "$1,246m (FY2015, strong USD hedge gains) to -$380m (FY2021, FX losses). "
    "The freeze removes currency translation and derivative fair-value movements "
    "from the model — a standard simplification.")

add_heading(doc, "4.8  Balance Sheet Supplemental Items", 2)
add_heading(doc, "4.8.1  Shares Outstanding (Rows 543–544)", 3)
add_table(doc,
    ["Year", "Period-end shares (m)", "Filing-date shares (m)", "Change"],
    [
        ("FY2015", "1,712", "1,710.7", "Buybacks ongoing"),
        ("FY2019", "1,568", "1,566.9", ""),
        ("FY2021", "1,578", "1,581.8", "Slight increase from SBC issuances"),
        ("FY2023", "1,531.9", "1,530.0", ""),
        ("FY2024", "1,503", "1,499.4", "Continued reduction; ~209m shares retired since FY2015"),
    ],
    [1.2, 2.0, 2.0, 1.5])
add_para(doc,
    "Row 543 = shares on filing date (actual); Row 544 = shares at fiscal year-end "
    "(rounded to millions). The small difference between them reflects buybacks or "
    "issuances between fiscal year-end (31 May) and the 10-K filing date (~July).")

add_heading(doc, "4.8.2  Net Debt History (Rows 548–549)", 3)
add_table(doc,
    ["Year", "Total Debt ($m)", "Cash+Investments ($m)", "Net Debt ($m)", "Status"],
    [
        ("FY2015", "1,260", "6,002", "-4,742", "Strong net cash"),
        ("FY2019", "3,853", "4,663", "-810", "Slight net cash"),
        ("FY2020", "13,015", "8,787", "+4,228", "FIRST net debt position — COVID borrowing"),
        ("FY2021", "12,813", "13,476", "-663", "Back to net cash (massive COO recovery)"),
        ("FY2022", "12,627", "12,997", "-370", "Marginal net cash"),
        ("FY2023", "12,144", "10,675", "+1,469", "Net debt again (buybacks exceeded FCF)"),
        ("FY2024", "11,983", "11,582", "+401", "Slight net debt"),
    ],
    [1.0, 1.8, 1.8, 1.0, 1.5])
add_para(doc,
    "NIKE crossed into net debt for only the second time in FY2023 and remained "
    "there in FY2024, driven by the $5.5bn FY2023 buyback program outpacing FCF. "
    "This is a meaningful shift for a company historically known for its net-cash "
    "fortress balance sheet.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  PART V — FREE CASH FLOW AND VALUATION MECHANICS
# ═══════════════════════════════════════════════════════════════
add_heading(doc, "PART V — Free Cash Flow and Valuation Mechanics", 1)

add_heading(doc, "5.1  EFCF — Equity Free Cash Flow", 2)
add_para(doc,
    "The model computes EFCF (Equity Free Cash Flow) as the FCF attributable to "
    "equity holders after debt service. The formula chain:")
add_para(doc,
    "    EFCF = NI + D&A − CapEx − ΔNOWC − Debt repayments + New debt issuances", bold=True)
add_para(doc,
    "This is the 'direct' approach to equity valuation (valuing equity directly "
    "rather than valuing the firm and subtracting debt). It is used alongside the "
    "WACC-based DCF as a cross-check.")

add_heading(doc, "5.2  CFADS — Cash Flow Available for Debt Service", 2)
add_para(doc,
    "CFADS (rows 260–261) measures pre-financing cash flow available to pay all "
    "debt obligations. Two alternative calculations are provided in the model:")
add_para(doc,
    "    Row 260: CFADS_A = EBIT(1−t) + D&A − CapEx − ΔNOWC\n"
    "    Row 261: CFADS_B = Net Income + D&A − CapEx − ΔNOWC + after-tax interest", bold=True)
add_callout(doc,
    "[CRITICAL BUG] In row 260, the tax rate reference is NOT anchored. "
    "Historical formula: C260: =C217+C229+C227*(1-R25)-C230-C231 (uses absolute $R$25). "
    "But D260: =D217+D229+D227*(1-S25)-D230-D231 (uses S25, not $R25). "
    "E260 uses T25, F260 uses U25, etc. These cells (S25, T25, U25...) are EMPTY "
    "in the spreadsheet, meaning the tax adjustment defaults to (1-0) = 1, i.e. "
    "0% tax. CFADS for FY2016 through FY2023 (columns D through K) are therefore "
    "computed with a ZERO tax rate rather than the ~21% statutory rate. "
    "The result: historical CFADS figures are OVERSTATED by approximately "
    "(EBIT × statutory rate) per year. FY2024 CFADS is correct (column L uses "
    "a formula without this bug). This systematic error affects 8 of 10 "
    "historical CFADS data points.", "CRITICAL BUG")

add_heading(doc, "5.3  Interest Expense (Row 356)", 2)
add_para(doc,
    "Interest expense = average IBD balance × average interest rate. "
    "The average IBD is computed as (beginning IBD + ending IBD) / 2, "
    "where IBD = sum of all rows annotated 'IBD' (rows 522, 523, 524, 528, 529). "
    "The average interest rate is drawn from a separate rate assumption cell "
    "(row 25 area) that is set to the approximate weighted average cost of NIKE's "
    "outstanding debt (~3.5% pre-tax). Cash interest paid (row 471 = $381m FY2024) "
    "confirms the accrual-basis interest expense in the model (~$382m FY2024) "
    "is accurate — less than $1m difference.")

add_heading(doc, "5.4  CapEx Derivation from Balance Sheet (Row 248)", 2)
add_para(doc,
    "Rather than pulling CapEx directly from the CF statement (row 447), the model "
    "DERIVES it from the BS changes in Gross PP&E:")
add_para(doc,
    "    CapEx = ending Gross PP&E − beginning Gross PP&E + asset disposals", bold=True)
add_para(doc,
    "This approach ensures the BS and CF statement are consistent. It also means "
    "that if revenue-scaling causes Gross PP&E to change, CapEx automatically adjusts. "
    "Asset disposals (row 448) are nil from FY2017 onwards, so in practice: "
    "CapEx = ΔGPPE. The model's forecast CapEx % of revenue (~3.2%) is consistent "
    "with NIKE's 10-year average (~3.0–3.5%).")

add_heading(doc, "5.5  ITS — Interest Tax Shield (Two Methods)", 2)
add_para(doc,
    "The model values the interest tax shield using two alternative approaches, "
    "each providing a check on the other:")
add_table(doc,
    ["Method", "Description", "Formula Row"],
    [
        ("Method 1: Modigliani-Miller (perpetuity)", "ITS PV = D × tax rate; assumes shield is permanent and valued at pre-tax cost of debt", "Row ~Q9 section"),
        ("Method 2: Miles-Ezzell (annual shield)", "Each year's shield = interest × tax rate, discounted at cost of unlevered equity; more conservative", "Row ~Q9 section"),
    ],
    [2.0, 4.5, 1.5])
add_para(doc,
    "The model uses Method 1 as the primary ITS for the APV calculation and "
    "Method 2 as a cross-check. Both are used to reconcile the levered and "
    "unlevered DCF values.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  PART VI — DCF VALUATION AND WACC
# ═══════════════════════════════════════════════════════════════
add_heading(doc, "PART VI — DCF Valuation and WACC", 1)

add_heading(doc, "6.1  WACC Calculation", 2)
add_para(doc,
    "The model uses a mixed WACC that blends the after-tax cost of debt and "
    "cost of equity using market-value weights. The D/(D+E) weight is derived from "
    "row 323 (D/TA ratio) which is converted to D/(D+E) using the identity: "
    "D/(D+E) = (D/TA) / (1 + D/TA − 1) = D/TA / (D/E + D/TA).")
add_para(doc,
    "WACC components (base case):\n"
    "  • Risk-free rate: 10-year US Treasury yield (~4.3%)\n"
    "  • Equity risk premium (ERP): ~5.5%\n"
    "  • Beta: ~0.8–0.9 (NIKE's 5-year monthly beta)\n"
    "  • Cost of equity = RFR + β × ERP ≈ 8.7–9.3%\n"
    "  • Pre-tax cost of debt ≈ 3.5–4.0%\n"
    "  • After-tax cost of debt ≈ 3.0% (at 21% tax rate)\n"
    "  • Base WACC ≈ 10.43% (cell O25)")

add_heading(doc, "6.2  Sensitivity Table — Exact WACC Values (Rows 63–74)", 2)
add_para(doc,
    "The sensitivity analysis creates a matrix with WACC on the vertical axis and "
    "terminal growth rate (g) on the horizontal axis. The exact WACC values are:")
add_table(doc,
    ["Row", "WACC Value", "Relationship"],
    [
        ("H63", "≈ 7.5%", "Fixed low end"),
        ("H64", "≈ 8.5%", "Fixed"),
        ("H65", "= H66 − 0.5%", "Derived relative to row below"),
        ("H66", "= H67 − 0.5%", "Derived relative to row below"),
        ("H67", "= 9.5%", "Fixed"),
        ("H68", "= O25 (base WACC)", "≈ 10.43%; the model's actual base case WACC"),
        ("H69", "= 11.5%", "Fixed high end"),
        ("H70", "= H69 + 0.5%", "= 12.0%"),
        ("H71", "= H70 + 1.0%", "= 13.0%"),
        ("H72", "= H71 + 1.0%", "= 14.0%"),
        ("H73", "= H72 + 1.0%", "= 15.0%"),
        ("H74", "= H73 + 1.0%", "= 16.0%"),
    ],
    [0.8, 1.5, 4.5])
add_para(doc,
    "Terminal growth rate (g) columns I through Q = 1%, 2%, 3%, 4% (base=M62), "
    "5%, 6%, 7%, 8%, 9%. The base g = M62 = 4% (cell M62 stores the base case "
    "long-term growth assumption). At WACC=10.43%, g=4%, the model produces "
    "its central DCF equity value.")

add_heading(doc, "6.3  IRR and NPV", 2)
add_para(doc,
    "The model also computes the Internal Rate of Return (IRR) implied by "
    "the projected FCF stream. The IRR is back-solved from the discounted "
    "cash flow series using Excel's IRR() function. This allows comparison "
    "of expected return against WACC: if IRR > WACC, value is being created.")

add_heading(doc, "6.4  Comparable Company Analysis (Sheet1)", 2)
add_para(doc,
    "Sheet1 contains CIQ data for a set of comparable companies (sportswear and "
    "consumer discretionary peers). The model does not appear to directly import "
    "multiples from Sheet1 into the valuation; rather, Sheet1 serves as a "
    "reference table for context. The 'valuatione' sheet has a comps section "
    "(rows ~60–100 approximately) with manually entered peer multiples.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  PART VII — PER SHARE AND DIVIDEND ANALYSIS
# ═══════════════════════════════════════════════════════════════
add_heading(doc, "PART VII — Per Share and Dividend Analysis", 1)

add_heading(doc, "7.1  EPS History (Row 363)", 2)
add_table(doc,
    ["Year", "Diluted EPS", "Notes"],
    [
        ("FY2015", "$1.85", "Pre-TCJA"),
        ("FY2016", "$2.16", ""),
        ("FY2017", "$2.51", ""),
        ("FY2018", "$1.17", "TCJA one-time charge; anomalously low"),
        ("FY2019", "$2.49", "Recovery"),
        ("FY2020", "$1.60", "COVID"),
        ("FY2021", "$3.56", "Post-COVID recovery"),
        ("FY2022", "$3.75", "Near peak"),
        ("FY2023", "$3.23", "Inventory write-down + restructuring"),
        ("FY2024", "$3.73", "Reported"),
        ("FY2024 normalised", "$2.92", "Strips restructuring + normalises tax"),
    ],
    [1.2, 1.5, 4.0])

add_heading(doc, "7.2  DPS History and CAGR (Row 374)", 2)
add_table(doc,
    ["Year", "DPS ($)", "YoY Growth"],
    [
        ("FY2015", "0.54", "—"),
        ("FY2016", "0.62", "+14.8%"),
        ("FY2017", "0.70", "+12.9%"),
        ("FY2018", "0.78", "+11.4%"),
        ("FY2019", "0.86", "+10.3%"),
        ("FY2020", "0.945", "+9.9%"),
        ("FY2021", "1.045", "+10.6%"),
        ("FY2022", "1.195", "+14.4%"),
        ("FY2023", "1.33", "+11.3%"),
        ("FY2024", "1.45", "+9.0%"),
    ],
    [1.2, 1.2, 2.0])
add_para(doc,
    "10-year DPS CAGR (FY2015–FY2024) = (1.45/0.54)^(1/9) - 1 ≈ 11.6%. "
    "NIKE has grown its dividend at double-digit rates for over a decade, "
    "qualifying as a Dividend Aristocrat candidate. However, the payout ratio "
    "has expanded significantly (see below), limiting future DPS growth unless "
    "EPS recovers substantially.")

add_heading(doc, "7.3  Payout Ratio History (Row 375)", 2)
add_table(doc,
    ["Year", "Payout Ratio", "Driver"],
    [
        ("FY2015", "29.2%", "Normal; growing NI"),
        ("FY2016", "28.7%", "Normal"),
        ("FY2017", "27.9%", "Normal"),
        ("FY2018", "66.7%", "Anomaly: TCJA one-off slashed EPS; DPS maintained"),
        ("FY2019", "34.5%", "Normalised"),
        ("FY2020", "59.1%", "COVID-year high payout"),
        ("FY2021", "29.3%", "NI recovery"),
        ("FY2022", "31.9%", "Normal"),
        ("FY2023", "41.2%", "NI weakness"),
        ("FY2024", "38.9%", "Forecast range"),
    ],
    [1.2, 1.5, 4.0])

add_heading(doc, "7.4  ADR Ratio (Row 376)", 2)
add_para(doc,
    "Row 376 stores ADR ratio = 0.1. NIKE Class B shares trade as ADRs at a "
    "1:10 ratio — each ADR represents 0.1 ordinary shares. This affects per-ADR "
    "EPS and DPS figures when comparing to non-US listed shares. "
    "For US-listed NKE (NYSE), shares trade directly (not as ADRs), "
    "so the ADR ratio is primarily relevant for international investors.")

add_heading(doc, "7.5  Book Value per Share (Row 545)", 2)
add_table(doc,
    ["Year", "BVPS ($)", "Notes"],
    [
        ("FY2015", "7.42", ""),
        ("FY2018", "6.13", "RE compression from TCJA"),
        ("FY2020", "5.14", "Lowest; COVID + negative RE"),
        ("FY2022", "9.73", "High NI year"),
        ("FY2024", "9.60", "Current"),
    ],
    [1.2, 1.5, 4.0])
add_para(doc,
    "Tangible BVPS (row 547) = $9.27 in FY2024, only slightly below total BVPS "
    "because goodwill and other intangibles are small ($499m combined = $0.33/share).")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  PART VIII — SUPPLEMENTAL DISCLOSURES
# ═══════════════════════════════════════════════════════════════
add_heading(doc, "PART VIII — Supplemental Disclosures", 1)

add_heading(doc, "8.1  Stock-Based Compensation Detail (Rows 396–412)", 2)
add_para(doc,
    "The SBC section contains more granularity than the headline SBC figure:")
add_heading(doc, "8.1.1  SBC Allocation History", 3)
add_table(doc,
    ["Row", "Category", "Available Years", "FY2024"],
    [
        ("396", "SBC in COGS", "FY2021 only ($4m)", "NA (not separately disclosed FY2024)"),
        ("397", "SBC in SG&A", "FY2015 ($191m) and FY2021 ($41m) only", "NA"),
        ("398", "SBC unallocated", "FY2016–FY2024", "$804m FY2024"),
        ("399", "Total SBC", "FY2015–FY2024", "$804m FY2024"),
        ("409", "Restricted stock comp (before tax)", "FY2015–FY2024", "$399m FY2024"),
        ("410", "Total SBC before tax", "FY2015–FY2024", "$804m FY2024 (same as row 399)"),
        ("411", "SBC tax effect", "FY2017–FY2022 only", "NA FY2024"),
        ("412", "SBC after tax", "FY2017–FY2022 only", "NA FY2024"),
    ],
    [0.5, 2.0, 2.2, 2.0])
add_para(doc,
    "Restricted stock compensation (row 409) = $399m vs total SBC = $804m; "
    "the remaining $405m = stock options and other equity awards. "
    "The mix has shifted dramatically: options were the dominant form pre-FY2018, "
    "but RSUs have become dominant post-FY2018 as NIKE aligned with S&P 500 norms.")
add_para(doc,
    "SBC tax effect (row 411) is only available for FY2017–FY2022. The after-tax "
    "SBC cost swung between -$12m (FY2018, TCJA year) and +$314m (FY2021). "
    "The negative FY2018 figure means NIKE's tax BENEFIT from SBC actually "
    "exceeded the gross SBC charge that year — the TCJA triggered a large "
    "final tax deduction on outstanding unvested awards.")

add_heading(doc, "8.1.2  Stock Options Activity (Rows 402–407)", 3)
add_table(doc,
    ["Row", "Item", "FY2015", "FY2016", "FY2017+"],
    [
        ("403", "Options outstanding beginning", "127.1m", "116.2m", "NA"),
        ("404", "Options granted", "18.4m", "20.6m", "NA"),
        ("405", "Options exercised", "27.2m", "22.5m", "NA"),
        ("406", "Options cancelled", "2.1m", "2.3m", "NA"),
        ("407", "Options outstanding ending", "116.2m", "112.0m", "NA"),
    ],
    [2.2, 1.3, 1.3, 2.0])
add_para(doc,
    "Stock option disclosures are only available for FY2015–FY2016. From FY2017 "
    "onwards, CIQ/NIKE ceased reporting a detailed options activity table separately "
    "(options data was rolled into the broader SBC footnote). The shift coincides "
    "with NIKE's transition from options-heavy to RSU-heavy compensation.")

add_heading(doc, "8.2  Tax Supplemental — NOL Carryforwards (Rows 563–572)", 2)
add_para(doc,
    "NIKE maintains foreign Net Operating Loss carryforwards (the US entity is "
    "consistently profitable). The NOL breakdown by expiration bucket:")
add_table(doc,
    ["Row", "Bucket", "FY2015 ($m)", "FY2020 ($m)", "FY2023 ($m)"],
    [
        ("564", "Current year", "NA", "NA", "NA"),
        ("565", "Year +1", "NA", "3", "NA"),
        ("566", "Year +2", "4", "2", "NA"),
        ("567", "Year +3", "1", "2", "NA"),
        ("568", "After 5 years", "17", "59", "61"),
        ("569", "No expiration", "14", "17", "NA"),
        ("570", "Total NOL", "36", "83", "61"),
    ],
    [1.8, 1.5, 1.2, 1.2, 1.2])
add_para(doc,
    "Total NOL peaked at $289m in FY2018 (TCJA transition year) then declined "
    "as profitable foreign subsidiaries absorbed the carryforwards. "
    "Row 572: Total Tax Benefit Carryforward = $208m in FY2017 only — "
    "this represents US R&D tax credits and other federal credits carried forward "
    "that exist on a single-year basis. All carryforward items are immaterial "
    "relative to NIKE's total deferred tax position.")

add_heading(doc, "8.3  FIN 48 / Uncertain Tax Positions (Rows 573–583)", 2)
add_para(doc,
    "FIN 48 (ASC 740-10) requires disclosure of unrecognised tax benefits — "
    "positions taken in tax returns that have not yet been sustained under audit. "
    "NIKE's UTB rollforward:")
add_table(doc,
    ["Year", "Beginning UTB", "Ending UTB", "Key Movement"],
    [
        ("FY2015", "$506m", "$438m", "-$123m reductions for prior years + -$27m settlements"),
        ("FY2018", "$461m", "$698m", "+$249m additions for current year — TCJA-related positions"),
        ("FY2020", "$808m", "$771m", "+$181m prior year additions, -$171m reductions"),
        ("FY2023", "$848m", "$936m", "+$95m prior year; modest growth"),
        ("FY2024", "$936m", "$990m", "+$77m current year; +$35m prior years; -$24m statute lapse; -$22m settlements"),
    ],
    [1.0, 1.3, 1.3, 3.2])
add_para(doc,
    "UTB has grown from $438m (FY2015) to $990m (FY2024) — a 126% increase — "
    "reflecting NIKE's increasingly complex international tax structure and "
    "ongoing disputes with tax authorities in multiple jurisdictions (primarily "
    "related to transfer pricing for intellectual property). "
    "Row 583: Interest and penalties recognised on the balance sheet (after-tax) "
    "= $332m in FY2024 (first year this line is disclosed by CIQ). This represents "
    "accrued interest on the UTB that NIKE expects to pay if challenged positions "
    "are resolved adversely.")

add_heading(doc, "8.4  Fair Value Measurements (Rows 585–590)", 2)
add_table(doc,
    ["Year", "Level 1 ($m)", "Level 2 ($m)", "Level 3 ($m)", "Total Assets ($m)"],
    [
        ("FY2015", "869", "6,229", "8", "7,106"),
        ("FY2020", "1,204", "7,082", "nil", "8,286"),
        ("FY2021", "2,892", "9,836", "nil", "12,728"),
        ("FY2022", "3,801", "9,237", "nil", "13,038"),
        ("FY2023", "2,655", "6,810", "nil", "9,465"),
        ("FY2024", "1,175", "9,528", "nil", "10,703"),
    ],
    [1.0, 1.2, 1.2, 1.2, 1.5])
add_para(doc,
    "Level 1 (quoted prices): primarily money market funds and publicly traded "
    "equity investments. FY2022 peak ($3,801m) reflects elevated cash in money "
    "market funds post-COVID before deployment into buybacks.\n"
    "Level 2 (observable but not quoted): corporate bonds, government securities, "
    "derivatives at FV. The $9,528m in FY2024 is predominantly NIKE's "
    "short-term investment portfolio and interest-rate / currency derivatives.\n"
    "Level 3 (unobservable inputs): only existed in FY2015–FY2017 ($8–10m), "
    "likely private equity interests since divested.")
add_para(doc,
    "Fair value of liabilities (row 590) = $7,782m FY2024. NIKE's debt is "
    "classified Level 2 (observable rates). Carrying value of IBD = "
    "$7,940m (LT $7,934m + current $6m ST). Market value ($7,782m) < carrying "
    "($7,940m) because market interest rates in 2024 are HIGHER than NIKE's "
    "coupon rates (most bonds were issued at 2–3%; current rates ~5–5.5%). "
    "The $158m premium (market value discount) would be a gain in a debt repurchase.")

add_heading(doc, "8.5  Operating Lease Adoption (ASC 842, FY2020)", 2)
add_para(doc,
    "NIKE adopted ASC 842 in FY2020 (fiscal year ended 31 May 2020). "
    "The impact was: $2,913m right-of-use asset added to long-term assets (row 517) "
    "and $445m current + $2,913m long-term lease liability added to the balance "
    "sheet (rows 524, 529). Total lease obligations = $3,358m on adoption date. "
    "Gross PP&E jumped from $9,469m (FY2019) to $12,758m (FY2020) — a $3.3bn "
    "increase — almost entirely from ROU assets being capitalised.")
add_para(doc,
    "CIQ row 551 'Debt Equivalent of Operating Leases' provides the CIQ-estimated "
    "PV of ALL operating lease commitments (including those beyond the on-BS period). "
    "FY2024 CIQ estimate = $8,408m vs on-BS $3,043m (current $477 + LT $2,566). "
    "The difference reflects: (1) CIQ using a different discount rate, "
    "(2) CIQ including renewal options not recognised on-BS.")

add_heading(doc, "8.6  Employee Headcount (Row 559)", 2)
add_table(doc,
    ["Year", "FT Employees", "Change"],
    [
        ("FY2015", "62,600", "Baseline"),
        ("FY2017", "74,400", "+18.9% in 2 years"),
        ("FY2020", "75,400", "COVID — minimal attrition"),
        ("FY2022", "79,100", "+5% post-COVID build"),
        ("FY2023", "83,700", "PEAK"),
        ("FY2024", "79,400", "-5.1%; ~4,300 roles eliminated; restructuring program"),
    ],
    [1.0, 1.5, 4.2])
add_para(doc,
    "The FY2024 headcount decline to 79,400 represents the first significant "
    "workforce reduction in NIKE's modern history. The restructuring program "
    "targets cost reduction and organisational simplification. The model does not "
    "separately forecast headcount-driven costs; SG&A% of revenue implicitly "
    "captures the cost structure going forward.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  PART IX — DUPOINT, RATIOS, AND ANALYTICS
# ═══════════════════════════════════════════════════════════════
add_heading(doc, "PART IX — DuPont Analysis and Financial Ratios", 1)

add_heading(doc, "9.1  DuPont Decomposition", 2)
add_para(doc,
    "The model includes a full DuPont decomposition section. NIKE's 5-factor "
    "DuPont analysis decomposes Return on Equity (ROE) as:")
add_para(doc,
    "    ROE = Net Margin × Asset Turnover × Financial Leverage\n"
    "    (or 5-factor: × Tax Burden × Interest Burden × EBIT Margin × Asset Turnover × Leverage)", bold=True)
add_table(doc,
    ["Year", "ROE", "Net Margin", "Asset Turnover", "Leverage", "Driver"],
    [
        ("FY2015", "~25%", "11.8%", "1.29x", "1.70x", "High turnover"),
        ("FY2019", "~44%", "10.3%", "1.65x", "2.62x", "Leverage increase"),
        ("FY2020", "~31%", "6.8%", "1.19x", "3.89x", "COVID: low NI, high leverage"),
        ("FY2022", "~44%", "12.9%", "1.16x", "2.64x", "High NI, good margins"),
        ("FY2024", "~40%", "11.1%", "1.35x", "2.64x", "Current"),
    ],
    [1.0, 0.8, 1.2, 1.5, 1.2, 2.0])
add_para(doc,
    "NIKE's ROE has been elevated by progressive leverage (total assets / equity) "
    "increasing from 1.7x (FY2015) to 2.6x (FY2024) as buybacks compressed equity. "
    "This financial engineering effect masks modest asset turnover growth.")

add_heading(doc, "9.2  Efficiency Ratios", 2)
add_table(doc,
    ["Ratio", "FY2015", "FY2024", "Formula"],
    [
        ("Days Inventory (DIO)", "~57 days", "~53 days", "Inventory / COGS × 365"),
        ("Days Sales Outstanding (DSO)", "~44 days", "~31 days", "AR / Revenue × 365; significant improvement"),
        ("Days Payable Outstanding (DPO)", "~28 days", "~36 days", "AP / COGS × 365"),
        ("Cash Conversion Cycle", "~73 days", "~48 days", "DIO + DSO − DPO; major improvement"),
        ("Revenue / Assets", "1.29x", "1.35x", "Asset turnover"),
    ],
    [1.8, 1.0, 1.0, 3.0])

add_heading(doc, "9.3  Liquidity Ratios", 2)
add_table(doc,
    ["Ratio", "FY2015", "FY2024", "Notes"],
    [
        ("Current Ratio", "2.46x", "2.40x", "Stable; strong current position"),
        ("Quick Ratio", "1.77x", "1.69x", "Excluding inventory"),
        ("Cash Ratio", "0.95x", "1.09x", "Cash+ST investments / current liabilities; >1 = highly liquid"),
    ],
    [1.8, 1.0, 1.0, 3.0])

add_heading(doc, "9.4  Leverage Ratios", 2)
add_table(doc,
    ["Ratio", "FY2015", "FY2020", "FY2024", "Notes"],
    [
        ("Total Debt / EBITDA", "0.17x", "1.72x", "1.59x", "Peaked FY2020; still manageable"),
        ("Net Debt / EBITDA", "-0.63x", "0.56x", "0.05x", "Near zero net leverage FY2024"),
        ("Total Assets / Equity", "1.70x", "3.89x", "2.64x", "Financial leverage; buyback-driven"),
        ("IBD / Total Assets", "5.8%", "41.5%", "31.4%", "D/TA from row 323 drivers"),
    ],
    [2.2, 0.9, 0.9, 0.9, 2.0])

add_heading(doc, "9.5  Coverage Ratios (Three Methods)", 2)
add_para(doc,
    "The model provides three different debt coverage metrics:")
add_table(doc,
    ["Coverage Metric", "FY2024 Value", "Formula", "Interpretation"],
    [
        ("Interest Coverage (EBIT basis)", "~18x", "EBIT / Interest expense", "Very strong; investment grade"),
        ("Debt Service Coverage (DSCR)", "~4x", "CFADS / (Interest + Principal)", "Measures ability to service all debt"),
        ("Fixed Charge Coverage (FCCR)", "~3.5x", "EBIT / (Interest + Leases + Sinking fund)", "Includes all fixed obligations"),
    ],
    [2.5, 1.3, 2.0, 2.0])
add_para(doc,
    "All three coverage ratios comfortably exceed typical investment-grade thresholds "
    "(DSCR > 1.25x, FCCR > 1.5x). NIKE's A1/AA- ratings are well-supported by "
    "these metrics.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  PART X — CHART DATA SECTIONS
# ═══════════════════════════════════════════════════════════════
add_heading(doc, "PART X — Chart Data Sections (Rows 108–165)", 1)

add_heading(doc, "10.1  Question 5 Chart Data (Rows 108–130)", 2)
add_para(doc,
    "Rows 108–130 contain pre-computed chart data series for the Q5 visualisation "
    "(a 17-year chart covering FY2015–FY2031). This section is separate from the "
    "core model and exists purely to feed charts with consistently-scaled values.")
add_table(doc,
    ["Row", "Series", "Unit Conversion", "Range Covered"],
    [
        ("~108", "Revenue (billions)", "÷ 10^6 (thousands to billions)", "FY2015–FY2031 all years"),
        ("~109", "COGS (billions)", "÷ 10^6", "FY2015–FY2031"),
        ("~110", "Net Income (billions)", "÷ 10^6", "FY2015–FY2031"),
        ("~115", "FCF (billions)", "÷ 10^6", "Forecast years only (FY2025–FY2031); historical rows blank"),
    ],
    [0.7, 2.0, 2.0, 2.5])
add_para(doc,
    "The conversion from the model's thousands-denomination to billions (dividing "
    "by 10^6) is done in this chart data section so the Y-axis labels are readable. "
    "Note that FCF is only charted for forecast years — this is because the "
    "historical FCF requires the Q1 definition which is not backfilled here.")

add_heading(doc, "10.2  Question 6 FCF Components Chart Data (Rows 133–165)", 2)
add_para(doc,
    "Rows 133–165 contain the FCF component breakdown for forecast years only, "
    "expressed in $100,000s (÷ 10^5 from the model's thousands base = ÷ 10^2 "
    "effectively). This allows the stacked bar chart for the FCF waterfall.")
add_table(doc,
    ["Row", "Component", "Sign Convention"],
    [
        ("~133", "Net Income", "Positive bar"),
        ("~137", "D&A", "Positive bar (add-back)"),
        ("~141", "Capital Expenditure", "Negative bar (cash outflow)"),
        ("~145", "Δ Net Working Capital", "Negative if WC increases"),
        ("~149", "After-tax Interest", "Positive (re-addition for unlevered FCF)"),
        ("~153", "Total FCF", "Sum; the final bar in waterfall"),
    ],
    [0.7, 2.5, 3.0])
add_para(doc,
    "Each row sources from the Q1 FCF build-up section (rows ~217–260). "
    "The ÷10^5 scaling makes the chart values in $100,000s (0.1 of a million), "
    "which displays as low single-digit numbers — unusual for a company of NIKE's "
    "size. Chart Y-axis would need labelling as '× $100,000' or '$0.1bn per unit'.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  PART XI — FORMULA BUGS AND MODEL QUIRKS
# ═══════════════════════════════════════════════════════════════
add_heading(doc, "PART XI — Formula Bugs and Model Quirks", 1)

add_heading(doc, "11.1  Confirmed Formula Bugs", 2)
add_callout(doc,
    "These are not modelling assumptions — they are formula errors that produce "
    "incorrect numerical results. Each has been verified by inspecting the raw "
    "cell formulas in fullsheet.txt.", "BUGS")

add_para(doc, "Bug 1: CFADS Non-Absolute Tax Rate Reference (Rows 260–261)", bold=True)
add_para(doc,
    "Severity: HIGH — affects 8 of 10 historical CFADS data points.\n"
    "Description: Column C (FY2015) uses absolute $R$25 for the tax rate; "
    "columns D through K (FY2016–FY2024) use relative references S25, T25... "
    "which are EMPTY cells. Result: historical CFADS uses 0% tax rate.\n"
    "Impact: CFADS overstated by approximately EBIT × 21% per year "
    "(≈ $1,100–$1,400m per year at current EBIT levels).\n"
    "Fix: Replace all relative rate references with $R$25 (or whichever cell "
    "stores the normalised tax rate).")

add_para(doc, "Bug 2: L220 Equity Raisings Formula Anomaly (Row 220)", bold=True)
add_para(doc,
    "Severity: LOW — numerical impact is ~$0.4m.\n"
    "Description: L220 (FY2025 equity raisings in the financing section) = "
    "L459 + L324. L324 = -0.4 (the payout ratio coefficient, a rate, not dollars). "
    "All other forecast cells (M220, N220, etc.) = just the corresponding row 459 "
    "value. Adding a fractional rate to a dollar flow is dimensionally incorrect.\n"
    "Fix: L220 should = L459 only (consistent with M220 = M459 etc).")

add_para(doc, "Bug 3: Long-term Debt One-Year Lag for FY2025 (Row 528)", bold=True)
add_para(doc,
    "Severity: LOW — only affects FY2025; FY2026+ are correct.\n"
    "Description: L528 = K518 × K323 (uses PRIOR year total assets). "
    "All subsequent years M528–R528 use current-year total assets. "
    "This makes FY2025 LT debt based on FY2024 assets, not FY2025.\n"
    "Fix: L528 should = L518 × L323.")

add_heading(doc, "11.2  Intentional Hard-Codes and Freezes", 2)
add_para(doc,
    "These are deliberate analyst choices, not bugs, but should be documented "
    "so future users don't mistake them for live formulas:")
add_table(doc,
    ["Row", "Item", "Hard-Coded Value", "Rationale"],
    [
        ("513 M–R", "Net PP&E", "$7,885m", "Comment 'Hard coded'; assumes maintenance capex only"),
        ("502 L–R", "Trading securities", "0", "NIKE exited trading portfolio; permanent zero"),
        ("509 L–R", "Other current assets", "$630m", "Frozen at FY2024; not revenue-driven"),
        ("522 L–R", "Short-term borrowings", "$6m", "Minimal commercial paper; assumed constant"),
        ("516 L–R", "Deferred tax assets LT", "$2,465m", "Complex to forecast; frozen"),
        ("514 L–R", "Goodwill", "$240m", "No acquisitions assumed"),
        ("538 L–R", "Accumulated OCI", "$53m", "No FX/hedge movements modelled"),
        ("530 L–R", "Pension obligations", "0", "NIKE's pension essentially defunded"),
        ("464 M–R", "Dividends (year 2–7)", "= L464 (frozen)", "Dividends frozen after FY2025"),
        ("468 L–R", "FX adjustment", "0", "USD-only model"),
        ("347 L–R", "Restructuring charges", "0", "Assumed completed by FY2024"),
    ],
    [1.2, 2.0, 1.3, 2.5])

add_heading(doc, "11.3  Unusual Formula Patterns", 2)
add_table(doc,
    ["Pattern", "Location", "Description"],
    [
        ("Explicit ^1 exponent", "Row 338 (revenue)", "L338: =K338*(1+L320)^1 — redundant but not harmful"),
        ("Year-on-year SG&A chaining", "Row 342", "M342 = L342/L338*M338 — scales off prior year ratio, not FY2024"),
        ("Dividend = buyback / 2", "Row 464", "L464: =L460/2 — approximation of historical ratio"),
        ("APIC includes buybacks", "Row 536", "L536: =K536+L460+L459 — L460 negative reduces APIC"),
        ("Row 357/358 hidden checks", "Income statement", "Unlabelled EBT/Tax and Tax@statutory rows for all historical years"),
        ("T530 inline WACC check", "Row 530, col T", "=K540/(K201+K540)*100 — equity weight % for WACC verification"),
        ("T533 standalone value", "Row 533, col T", "54.63; appears to be a liabilities/capital ratio hand-check"),
        ("S513 text comment", "Net PP&E row", "S513: 'Hard coded' — documents the freeze"),
        ("S522–S529 IBD labels", "Liability rows", "'IBD' text flags interest-bearing debt for WACC/coverage calcs"),
    ],
    [2.0, 1.5, 3.5])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  PART XII — COMPLETE HISTORICAL DATA REFERENCE
# ═══════════════════════════════════════════════════════════════
add_heading(doc, "PART XII — Complete Historical Data Reference", 1)

add_heading(doc, "12.1  Income Statement Summary (All Historical Years)", 2)
add_table(doc,
    ["Item", "FY15", "FY16", "FY17", "FY18", "FY19", "FY20", "FY21", "FY22", "FY23", "FY24"],
    [
        ("Revenue $m", "27,790", "32,376", "34,350", "36,397", "39,117", "37,403", "44,538", "46,710", "51,217", "51,362"),
        ("Gross Margin %", "46.2%", "44.4%", "44.1%", "43.8%", "44.7%", "43.4%", "44.8%", "44.3%", "44.0%", "43.6%"),
        ("SG&A $m", "8,834", "10,469", "10,563", "10,563", "11,630", "12,485", "12,275", "13,046", "14,685", "15,829"),
        ("EBIT $m", "3,980", "3,916", "4,199", "4,445", "4,771", "3,913", "5,771", "6,421", "5,931", "6,754"),
        ("EBIT margin %", "14.3%", "12.1%", "12.2%", "12.2%", "12.2%", "10.5%", "13.0%", "13.7%", "11.6%", "13.1%"),
        ("Net Income $m", "3,273", "3,760", "4,240", "1,933", "4,029", "2,539", "5,727", "6,046", "5,070", "5,700"),
        ("EPS diluted $", "1.85", "2.16", "2.51", "1.17", "2.49", "1.60", "3.56", "3.75", "3.23", "3.73"),
        ("EBITDA $m", "4,629", "4,578", "4,915", "5,219", "5,491", "5,032", "6,568", "7,261", "6,634", "7,550"),
    ],
    [2.0, 0.65, 0.65, 0.65, 0.65, 0.65, 0.65, 0.65, 0.65, 0.65, 0.65])

add_heading(doc, "12.2  Cash Flow Statement Summary", 2)
add_table(doc,
    ["Item", "FY15", "FY20", "FY21", "FY22", "FY23", "FY24"],
    [
        ("CFO $m", "4,680", "2,485", "6,657", "5,188", "5,841", "7,429"),
        ("CapEx $m", "-963", "-1,086", "-695", "-758", "-969", "-812"),
        ("CFF $m", "-2,790", "+2,491", "-1,459", "-4,836", "-7,447", "-5,888"),
        ("Net ΔCash $m", "+1,632", "+3,882", "+1,541", "-1,315", "-1,133", "+2,419"),
        ("Cash interest paid $m", "53", "140", "293", "290", "347", "381"),
        ("Cash taxes paid $m", "1,262", "1,028", "1,177", "1,231", "1,517", "1,299"),
        ("Total buybacks $m", "-2,534", "-3,067", "-608", "-4,014", "-5,480", "-4,250"),
        ("Total dividends $m", "-899", "-1,452", "-1,638", "-1,837", "-2,012", "-2,169"),
    ],
    [2.0, 0.8, 0.8, 0.8, 0.8, 0.8, 0.8])

add_heading(doc, "12.3  Balance Sheet Summary", 2)
add_table(doc,
    ["Item", "FY15", "FY19", "FY20", "FY22", "FY23", "FY24"],
    [
        ("Cash+ST Investments $m", "6,002", "4,663", "8,787", "12,997", "10,675", "11,582"),
        ("Total Current Assets $m", "15,587", "16,525", "20,556", "28,213", "25,202", "25,382"),
        ("Net PP&E $m", "3,011", "4,744", "7,963", "7,717", "8,004", "7,718"),
        ("Total Assets $m", "21,597", "23,717", "31,342", "40,321", "37,531", "38,110"),
        ("Total Current Liabilities $m", "6,332", "7,866", "8,284", "10,730", "9,256", "10,593"),
        ("Long-term Debt $m", "1,079", "3,464", "9,406", "8,920", "8,927", "7,934"),
        ("Long-term Leases $m", "nil", "342", "2,913", "2,777", "2,786", "2,566"),
        ("Total Equity $m", "12,707", "9,040", "8,055", "15,281", "14,004", "14,430"),
        ("Net Debt $m", "-4,742", "-810", "+4,228", "-370", "+1,469", "+401"),
    ],
    [2.0, 0.8, 0.8, 0.8, 0.8, 0.8, 0.8])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  PART XIII — CAPITAL ALLOCATION HISTORY
# ═══════════════════════════════════════════════════════════════
add_heading(doc, "PART XIII — Capital Allocation History", 1)

add_heading(doc, "13.1  Annual Shareholder Returns", 2)
add_table(doc,
    ["Year", "CFO $m", "CapEx $m", "FCF $m", "Buybacks $m", "Dividends $m", "Total Return $m", "Return/FCF %"],
    [
        ("FY2015", "4,680", "-963", "3,717", "2,534", "899", "3,433", "92%"),
        ("FY2016", "3,399", "-1,143", "2,256", "3,260", "1,022", "4,282", "190%"),
        ("FY2017", "3,846", "-1,105", "2,741", "3,252", "1,133", "4,385", "160%"),
        ("FY2018", "4,955", "-1,028", "3,927", "4,254", "1,243", "5,497", "140%"),
        ("FY2019", "5,903", "-1,119", "4,784", "4,286", "1,332", "5,618", "117%"),
        ("FY2020", "2,485", "-1,086", "1,399", "3,067", "1,452", "4,519", "323%"),
        ("FY2021", "6,657", "-695", "5,962", "608", "1,638", "2,246", "38%"),
        ("FY2022", "5,188", "-758", "4,430", "4,014", "1,837", "5,851", "132%"),
        ("FY2023", "5,841", "-969", "4,872", "5,480", "2,012", "7,492", "154%"),
        ("FY2024", "7,429", "-812", "6,617", "4,250", "2,169", "6,419", "97%"),
    ],
    [0.8, 0.7, 0.7, 0.8, 1.0, 1.2, 1.2])
add_para(doc,
    "NIKE consistently returns 90–160% of FCF to shareholders in most years. "
    "FY2020 and FY2016/FY2017 show >100% returns, funded by balance sheet borrowing. "
    "FY2021 was the exception: post-COVID buybacks were deliberately reduced "
    "(only $608m) as NIKE rebuilt inventory and supply chain flexibility. "
    "Cumulative FY2015–FY2024 capital returned: ~$53.6bn ($33.0bn buybacks + $14.7bn dividends).")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  APPENDIX A — COMPLETE FORMULA INVENTORY
# ═══════════════════════════════════════════════════════════════
add_heading(doc, "APPENDIX A — Critical Formula Reference", 1)

add_heading(doc, "A.1  Key Income Statement Formulas", 2)
add_table(doc,
    ["Row", "Cell (FY2025)", "Formula", "What It Computes"],
    [
        ("338", "L338", "=K338*(1+L320)^1", "Revenue = prior year × (1+growth)"),
        ("339", "L339", "=L338*L321", "COGS = Revenue × COGS%"),
        ("340", "L340", "=L338-L339", "Gross Profit"),
        ("342", "L342", "=(K342/K338)*L338", "SG&A = prior year SG&A% × current revenue"),
        ("347", "L347", "=0", "Restructuring (frozen)"),
        ("350", "L350", "=L340-L342-L347", "EBIT"),
        ("356", "L356", "=-avgIBD*rate", "Interest expense"),
        ("355", "L355", "=L350+int_inc-L356", "EBT"),
        ("359", "L359", "=L355*taxRate", "Tax provision"),
        ("361", "L361", "=L355-L359", "Net Income"),
    ],
    [0.5, 0.9, 2.5, 2.5])

add_heading(doc, "A.2  Key Balance Sheet Forecast Formulas", 2)
add_table(doc,
    ["Row", "Cell (FY2025)", "Formula", "What It Computes"],
    [
        ("500", "L500", "=K500+L236", "Cash = prior + ΔCash from model"),
        ("501", "L501", "=K501/K338*L338", "ST investments revenue-scaled"),
        ("504", "L504", "=K504/K338*L338", "Accounts receivable revenue-scaled"),
        ("507", "L507", "=K507/K338*L338", "Inventory revenue-scaled"),
        ("513", "L513", "=K513/K338*L338 (M+: hardcode)", "Net PP&E: rev-scaled FY25 then fixed"),
        ("514", "L514", "=K514", "Goodwill frozen"),
        ("516", "L516", "=K516", "DTA frozen"),
        ("528", "L528", "=K518*K323 (BUG: should be L518)", "LT debt = assets × D/TA ratio"),
        ("536", "L536", "=K536+L460+L459", "APIC += buybacks + issuances"),
        ("537", "L537", "=L361+L464+K537", "RE += NI + dividends + prior RE"),
    ],
    [0.5, 0.9, 2.5, 2.5])

add_heading(doc, "A.3  Key Cash Flow Forecast Formulas", 2)
add_table(doc,
    ["Row", "Cell (FY2025)", "Formula", "What It Computes"],
    [
        ("459", "L459", "=K459", "SBC proceeds frozen at FY2024"),
        ("460", "L460", "=L361*L324", "Buybacks = NI × payout_ratio (negative)"),
        ("464", "L464", "=L460/2", "Dividends = buybacks/2"),
        ("464", "M464", "=L464", "Dividends frozen from FY2026 onwards"),
        ("445", "L445", "=L233", "CF from Ops link to model"),
        ("451", "L451", "=L234", "CF from Investing link to model"),
        ("466", "L466", "=L235", "CF from Financing link to model"),
        ("469", "L469", "=L445+L451+L466+L468", "Net ΔCash (L468=0)"),
    ],
    [0.5, 0.9, 2.5, 2.5])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  APPENDIX B — MODELLING ASSUMPTIONS SUMMARY
# ═══════════════════════════════════════════════════════════════
add_heading(doc, "APPENDIX B — Modelling Assumptions and Limitations", 1)

add_heading(doc, "B.1  Key Assumptions", 2)
add_table(doc,
    ["Assumption", "Value / Approach", "Comment"],
    [
        ("Revenue growth rate (FY2025+)", "Driven by row 320; ~5-7% base case", "No geographic or segment detail"),
        ("COGS % of revenue", "~56.4% (FY2024 level)", "No margin expansion modelled"),
        ("SG&A % of revenue", "~30.8% (FY2024 level)", "Including marketing and D&A"),
        ("Restructuring charges", "Zero from FY2025", "May be optimistic if restructuring continues"),
        ("CapEx / Revenue", "~3.2%", "Below historical peak; assumes mature investment cycle"),
        ("D&A / Revenue", "~2.0%", "Consistent with recent years"),
        ("Tax rate", "~18% (post-TCJA normalised)", "Excludes SBC windfall effects"),
        ("Buyback payout ratio", "~40% of NI (row 324 = -0.4)", "Historical average; could be revised down"),
        ("Dividend policy", "= buybacks/2 in FY2025, frozen thereafter", "Oversimplification; dividends should grow"),
        ("WACC", "~10.43%", "Base case; see sensitivity table (Section 6.2)"),
        ("Terminal growth rate", "~4.0%", "Aggressive; equals long-term nominal GDP growth"),
        ("Net PP&E", "Frozen at $7,885m from FY2026", "Assumes maintenance capex replaces all D&A"),
        ("Goodwill", "Frozen at $240m", "No acquisitions"),
        ("FX impact", "Zero", "USD-only model"),
        ("Lease obligations", "Revenue-scaled", "Proportional to store count / revenue"),
    ],
    [2.5, 1.8, 2.5])

add_heading(doc, "B.2  Key Limitations", 2)
add_para(doc,
    "1. NO SEGMENT ANALYSIS: NIKE reports four geographic segments (North America, "
    "EMEA, Greater China, APAC) plus NIKE Direct vs wholesale. The model aggregates "
    "these; margin differences across segments are lost.\n\n"
    "2. NO PRODUCT MIX: Footwear (55% of revenue, ~45% margin) vs apparel (30%, "
    "~35% margin) vs equipment (5%+). Mix shift toward DTC/direct has margin "
    "implications not captured.\n\n"
    "3. CONSTANT MARGINS: Holding COGS% and SG&A% flat ignores NIKE's publicised "
    "margin recovery program (target: gross margin back to 46%+). If achieved, "
    "NPV would be materially higher.\n\n"
    "4. SINGLE WACC: The model uses one WACC for the entire forecast period, ignoring "
    "potential capital structure changes as NIKE repays debt.\n\n"
    "5. CFADS HISTORICAL BUG: 8 of 10 historical CFADS years are incorrect (zero tax) "
    "due to non-absolute cell references (see Section 5.2). Historical coverage "
    "analysis should not use these figures.\n\n"
    "6. SENSITIVITY GRID COVERAGE: WACC range 7.5%–16% and g range 1%–9% is very "
    "wide; the economic g range for a mature company is 0%–5%. The upper scenarios "
    "are not realistic.\n\n"
    "7. NORMALISED EPS NOT USED IN DCF: The model forecasts NI without "
    "restructuring charges (frozen to zero) which means forecast NI is already "
    "effectively 'normalised' — but FY2024 base year NI includes $443m restructuring "
    "which suppresses the base. This creates a built-in growth effect in FY2025 "
    "that is partly mechanical rather than operational.")

add_heading(doc, "B.3  Data Quality Notes", 2)
add_para(doc,
    "  • FY2016 and FY2017 CF statement data: Marked 'RS' (restated) by CIQ. "
    "The restatements relate to NIKE's retrospective adoption of ASU 2016-09 "
    "(SBC tax benefits reclassified from financing to operating activities). "
    "The restated figures show higher historical CFO than originally reported.\n\n"
    "  • FY2015–FY2017 balance sheet data: Also marked 'RS' or 'RUP' (rolled-up). "
    "Some items may be regrouped vs original 10-K filings.\n\n"
    "  • Stock option data availability: Only FY2015–FY2016 (rows 403–407). "
    "Post-FY2017 options data is not separately disclosed by CIQ.\n\n"
    "  • SBC tax effect: Only available FY2017–FY2022 (rows 411–412).\n\n"
    "  • FIN 48 interest and penalties: Only disclosed for FY2024 ($332m).\n\n"
    "  • Inventory breakdown (finished goods vs other): Only available "
    "FY2021–FY2023 in the CIQ dataset (rows 552–553).")

# ═══════════════════════════════════════════════════════════════
#  APPENDIX C — GLOSSARY
# ═══════════════════════════════════════════════════════════════
add_heading(doc, "APPENDIX C — Glossary", 1)
add_table(doc,
    ["Term", "Definition"],
    [
        ("APIC", "Additional Paid-In Capital; equity contributed above par value from share issuances and SBC exercises"),
        ("APV", "Adjusted Present Value; firm value = unlevered value + PV of tax shields"),
        ("ASC 842", "US GAAP lease accounting standard effective 2019; requires operating leases on balance sheet"),
        ("CFADS", "Cash Flow Available for Debt Service; pre-financing, pre-tax-shield measure of debt capacity"),
        ("CIQ", "S&P Capital IQ; the financial data provider whose raw export forms the model's historical data"),
        ("CIP", "Construction in Progress; PP&E not yet placed in service"),
        ("DIO", "Days Inventory Outstanding; inventory / COGS × 365"),
        ("DPO", "Days Payable Outstanding; AP / COGS × 365"),
        ("DSO", "Days Sales Outstanding; AR / Revenue × 365"),
        ("DTA/DTL", "Deferred Tax Asset / Deferred Tax Liability; timing differences between tax and book accounting"),
        ("EFCF", "Equity Free Cash Flow; FCF attributable to equity holders after debt service"),
        ("EBITDAR", "EBITDA before Rent; adds back operating lease costs; relevant for retail comparisons"),
        ("ETR", "Effective Tax Rate; income tax provision / pre-tax income"),
        ("FIN 48 / ASC 740-10", "US GAAP standard requiring disclosure of uncertain tax positions (UTBs)"),
        ("FY", "Fiscal Year; NIKE's FY ends 31 May each year (e.g. FY2024 = June 2023 to May 2024)"),
        ("IBD", "Interest-Bearing Debt; rows annotated 'IBD' in column S of the model"),
        ("IRR", "Internal Rate of Return; discount rate that makes NPV of FCF = 0"),
        ("ITS", "Interest Tax Shield; PV of tax savings from debt deductions"),
        ("LFCF", "Levered Free Cash Flow; CFO − CapEx (after interest)"),
        ("NOL", "Net Operating Loss carryforward; prior year tax losses used to offset future taxable income"),
        ("NOPAT", "Net Operating Profit After Tax; EBIT × (1 − tax rate)"),
        ("NOWC", "Net Operating Working Capital; current operating assets minus current operating liabilities"),
        ("OCI/AOCI", "Other Comprehensive Income / Accumulated OCI; items bypassing income statement (FX, hedges)"),
        ("PP&E", "Property, Plant & Equipment"),
        ("ROU", "Right-of-Use asset; on-balance-sheet lease asset under ASC 842"),
        ("RSU", "Restricted Stock Unit; equity award that vests on schedule"),
        ("SBC", "Stock-Based Compensation; non-cash expense for equity grants to employees"),
        ("TCJA", "Tax Cuts and Jobs Act of 2017; reduced US corporate rate from 35% to 21%"),
        ("ULFCF", "Unlevered Free Cash Flow; FCF before interest payments (pre-financing)"),
        ("UTB", "Unrecognised Tax Benefit; tax position taken in a return but not sustained under audit (FIN 48)"),
        ("WACC", "Weighted Average Cost of Capital; blended required return on debt + equity"),
    ],
    [1.5, 5.5])

# ─────────────────────────────────────────────────────────────
#  SAVE
# ─────────────────────────────────────────────────────────────
out_path = r"c:\Users\Rohan\Downloads\New folder (4)\NIKE Valuation - Complete Analysis.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
