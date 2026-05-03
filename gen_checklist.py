"""
gen_checklist.py — Append PART 80: Master Implementation Checklist to Architecture Plan.
Organized by build phase. Each item maps to specific Parts in the document.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCX_PATH = r'c:/Users/Rohan/Downloads/New folder (4)/Automated Valuation System - Architecture Plan.docx'

doc = Document(DOCX_PATH)

def h1(text):
    p = doc.add_paragraph(text)
    p.style = doc.styles['Heading 1']
    return p

def h2(text):
    p = doc.add_paragraph(text)
    p.style = doc.styles['Heading 2']
    return p

def h3(text):
    p = doc.add_paragraph(text)
    p.style = doc.styles['Heading 3']
    return p

def body(text):
    return doc.add_paragraph(text)

def item(text, indent=0):
    """Checkbox list item."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25 * (indent + 1))
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(f'\u2610  {text}')   # ballot box ☐
    run.font.size = Pt(10)
    return p

def done(text, indent=0):
    """Pre-ticked item (already architecturally designed — not yet coded)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25 * (indent + 1))
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(f'\u2610  {text}')
    run.font.size = Pt(10)
    return p

def ref(text):
    """Small reference note."""
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(8.5)
        run.font.italic = True
    return p

# ─────────────────────────────────────────────────────────────────────────────
doc.add_page_break()

h1('PART 80 — MASTER IMPLEMENTATION CHECKLIST')
body(
    'This checklist tracks every deliverable required to build the Automated Valuation '
    'System from scratch. Each item maps to the Part(s) where the specification lives. '
    'Work through phases in order — later phases depend on earlier ones. '
    'Mark each box as you complete it.'
)
body(
    'LEGEND:  ☐ = Not started / in progress.  '
    'Tick the box when the module is coded, tested, and integrated into main.py.'
)

# ═══════════════════════════════════════════════════════════════════════════
h2('PHASE 0 — Project Scaffold & Environment')
ref('Parts: 9, 11, 25, 33, 39, 49, 56')
# ═══════════════════════════════════════════════════════════════════════════

item('Create Python package: auto_valuation/ directory with __init__.py')
item('Create requirements.txt (pinned versions: yfinance, requests, pandas, numpy, openpyxl, scipy, python-dotenv)')
ref('Part 25, Part 9.2')
item('Create .env.example and .gitignore (never commit .env or API keys)')
ref('Part 39.2')
item('Create .env file locally with FMP_API_KEY, FRED_API_KEY')
item('Create config.py — all constants, 4-layer config hierarchy (global → sector → ticker-override → CLI)')
ref('Parts 33.1, 66')
item('Create overrides/ directory and overrides/EXAMPLE.json with full schema')
ref('Parts 31, 66.2')
item('Create main.py with argparse CLI: --ticker, --exchange, --currency, --scenario, --override, --batch')
ref('Parts 9.3, 11, 47.3')
item('Create logs/ directory; implement logging.py with audit trail (timestamp, ticker, version, warnings)')
ref('Part 33.3')
item('Write README.md with setup instructions, required API keys, usage examples')
ref('Part 56.1')
item('Write CHANGELOG.md (start at v1.0.0)')
ref('Part 56.2')
item('Set up tests/ directory with conftest.py and test data fixtures')
ref('Part 49.2')

# ═══════════════════════════════════════════════════════════════════════════
h2('PHASE 1 — Data Layer: Fetching & Cleaning')
ref('Parts: 2, 8, 13, 28, 29, 34, 39, 40, 46, 55, 60, 61')
# ═══════════════════════════════════════════════════════════════════════════

h3('1A — FMP Endpoints')
item('Implement fetch_income_statement() — FMP /v3/income-statement/{TICKER}?limit=10&apikey=')
ref('Part 2.1, A.1')
item('Implement fetch_balance_sheet() — FMP /v3/balance-sheet-statement/{TICKER}?limit=10&apikey=')
item('Implement fetch_cash_flow() — FMP /v3/cash-flow-statement/{TICKER}?limit=10&apikey=')
item('Implement fetch_quarterly_financials() — same endpoints with period=quarter for TTM computation')
ref('Part 28.1')
item('Implement fetch_profile() — FMP /v3/profile/{TICKER} (sector, industry, currency, exchange, shares)')
item('Implement fetch_income_quarterly(), fetch_balance_quarterly(), fetch_cashflow_quarterly() for TTM')
ref('Part 28')
item('Implement fetch_ntm_estimates() — FMP /v3/analyst-estimates/{TICKER} for forward consensus')
ref('Part 37.1, Appendix A')
item('Implement fetch_segment_data() — FMP /v4/revenue-product-segmentation and /v4/revenue-geographic-segmentation')
ref('Part 45.1')
item('Implement fetch_sec_filings_8k() — FMP /v3/sec_filings/{TICKER}?type=8-K (for pro forma warnings)')
ref('Part 78.2')

h3('1B — yfinance Data')
item('Implement fetch_yfinance_info() — .info dict: marketCap, currentPrice, fiftyTwoWeekHigh, fiftyTwoWeekLow, beta')
ref('Parts 79.2, A.2')
item('Implement fetch_52wk_range() — fiftyTwoWeekLow / fiftyTwoWeekHigh with 1-year price history fallback')
ref('Part 79.2')
item('Implement check_price_freshness() — warn if price data older than 3 days')
ref('Part 55.1')

h3('1C — FRED / Macro Data')
item('Implement fetch_risk_free_rate() — FRED 10-year Treasury yield (GS10 series) via FRED API')
ref('Parts 4.3, A.3')
item('Implement fetch_gdp_growth() — FRED nominal GDP growth for terminal growth ceiling check')
ref('Parts 3.3, A.3')

h3('1D — Damodaran Static Data')
item('Download and cache Damodaran industry beta files (updated every January)')
ref('Part A.4, 4.3')
item('Download and cache Damodaran ERP, size premium, CRP tables')
ref('Parts 38, A.4')

h3('1E — Data Cleaning & Normalisation')
item('Implement unit_normalize() — detect FMP "in thousands / millions / billions" and standardise to millions')
ref('Part 2.4')
item('Implement deduplicate_financial_data() — remove restated/duplicate periods (same fiscal year-end)')
ref('Part 46.1')
item('Implement standardise_field_names() — FMP snake_case → internal canonical field names')
ref('Part 2.8')
item('Implement compute_ttm() — sum last 4 quarters for income statement and cash flow; use latest quarter balance sheet')
ref('Part 28')
item('Implement apply_fx_conversion() — historical average rate for IS/CF, closing rate for BS, forecast in reporting currency')
ref('Part 29')
item('Implement align_fiscal_year() — convert non-Dec fiscal year to calendar year or pad to same-year stub')
ref('Parts 2.3, 4.5')
item('Implement calendarize_peer_data() — stub-period adjustment so all peers share same fiscal year end for comps')
ref('Part 37.2')
item('Implement detect_ma_year() — flag years with >15% revenue jump (acquisition noise)')
ref('Part 2.7')
item('Implement strip_discontinued_ops() — remove discontinued items from EBIT, UFCF, and balance sheet')
ref('Part 42.1')
item('Implement normalize_one_time_items() — add back goodwill impairment, restructuring, legal settlements')
ref('Parts 2.6, 42.2')

h3('1F — Validation at Data Layer')
item('Implement validate_fmp_data() — check for None/NaN in all critical fields; halt or warn per field criticality')
ref('Part 61.1')
item('Implement check_revenue_sanity() — halt if revenue < 0 or YoY growth > 200%')
ref('Part 61.2')
item('Implement check_nowc_sign() — negative NWC is valid (Amazon/Costco pattern); do NOT flag as error')
ref('Part 40.1')
item('Implement canonical compute_net_debt() — ST_debt + LT_debt + finance_leases - cash - ST_investments + preferred + NCI')
ref('Part 60')

# ═══════════════════════════════════════════════════════════════════════════
h2('PHASE 2 — Core Income Statement & UFCF Computation')
ref('Parts: 3, 13, 14, 15, 17, 19, 20, 23, 43, 44, 50, 51, 53, 55, 57, 58, 67, 68')
# ═══════════════════════════════════════════════════════════════════════════

h3('2A — Revenue Forecasting')
item('Implement compute_revenue_forecast() — mean-reversion from LTM growth toward sector median (Part 4.1/A.4 formula)')
ref('Parts 4.1, A.4')
item('Implement compute_revenue_bridge() — price × volume × mix decomposition (historical decomposition)')
ref('Part 57')
item('Implement compute_segment_forecast() — segment-level revenue drivers, sum to consolidated total')
ref('Part 45.2')
item('Implement check_revenue_recognition_flags() — AR days acceleration, deferred revenue spikes (ASC 606 red flags)')
ref('Part 55.3')

h3('2B — EBIT / EBITA')
item('Implement compute_ebit_margin_schedule() — fade EBIT margin from current toward sector median over forecast period')
ref('Part 51.1')
item('Implement normalize_effective_tax_rate() — cap at statutory rate, floor at 5%, 3-year average')
ref('Part 43.1')
item('Implement compute_nopat() = EBIT × (1 - effective_tax_rate)')
ref('Parts 3.1, 52.1')
item('Implement compute_nopat_nci_adjusted() — subtract minority share of NOPAT when NCI is material (>5% of NOPAT)')
ref('Part 67')
item('Handle non-operating income / expense below EBIT (recurring vs non-recurring filter)')
ref('Part 43.2, 64')
item('Implement compute_other_income_forecast() — scale recurring other income as % of revenue')
ref('Part 64')
item('Implement compute_pension_expense_forecast() — service cost + interest cost; check_pension_materiality()')
ref('Part 68')
item('Implement R&D capitalization (optional, off by default) — adjust EBIT and CapEx for tech/pharma')
ref('Part 55.2')

h3('2C — Working Capital')
item('Implement compute_working_capital_days() — AR days, AP days, Inventory days using average balances')
ref('Parts 4.2, 32.2')
item('Implement forecast_nowc() — project WC balances from days; clip negative NOWC (valid for retailer pattern)')
ref('Parts 4.2, 15.2, 40.1')
item('Implement check_wc_seasonality_flag() — detect high seasonal swing (>30% rev difference Q1 vs Q3)')
ref('Part 48.2')

h3('2D — Capital Expenditure')
item('Implement compute_capex_forecast() — split maintenance capex (% of PP&E) vs growth capex (% of rev growth)')
ref('Parts 51.2, 32.3')
item('Implement capex_convergence_to_da() — in terminal year, enforce capex → depreciation (steady-state constraint)')
ref('Part 51.2')
item('Handle asset-light CapEx anomaly — if CapEx reported < 0 (FMP sign error), take absolute value')
ref('Part 40.2')

h3('2E — Depreciation & Non-Cash Items')
item('Implement rollforward_ppe() — opening PP&E + CapEx - Depreciation = closing PP&E')
ref('Part 3.7')
item('Implement compute_da_forecast() — depreciation as % of opening PP&E; amortization as % of intangibles')
ref('Parts 3.1, 3.7')
item('Implement deferred_tax_rollforward() — DTA/DTL rollforward; include change in deferred tax in UFCF')
ref('Parts 53.1, 3.1')
item('Implement goodwill_rollforward() — opening goodwill + acquisitions - impairment = closing')
ref('Part 53.2')
item('Implement intangibles_amortization_rollforward()')
ref('Part 53.3')

h3('2F — UFCF Assembly')
item('Implement compute_ufcf() — NOPAT + D&A + change_in_deferred_tax - delta_NOWC - CapEx')
ref('Parts 3.1, 20, 23')
item('Verify SBC is excluded from UFCF (SBC is non-cash but reduces UFCF via dilution — include in diluted shares only)')
ref('Part A.1, 41.2')
item('Implement compute_historical_ufcf() — validate formula against FMP historical data for back-test')
ref('Part 32.3')
item('Add UFCF row to model sheet with LTM anchor + 5-year forecast')

h3('2G — Debt, Interest & Equity Rollforwards')
item('Implement rollforward_debt_schedule() — opening IBD + new_draws - repayments = closing IBD')
ref('Part 58.1')
item('Implement compute_interest_expense() — closing_IBD × avg_cost_of_debt; iterative solver for circularity')
ref('Parts 3.5 correction, C.2, A.3')
item('Implement compute_interest_income() — opening_cash × interest_income_rate (% of revenue method)')
ref('Part 19')
item('Implement rollforward_retained_earnings() — opening RE + net_income - dividends = closing RE')
ref('Part 76')
item('Implement rollforward_apic() — opening APIC + SBC + equity_issuances = closing APIC')
ref('Part 14.1')
item('Implement rollforward_aoci() — opening AOCI + OCI_items = closing AOCI')
ref('Part 14.3')
item('Implement check_balance_sheet_closes() — assert assets = liabilities + equity within $1m tolerance')
ref('Part 76')

h3('2H — Share Count')
item('Implement compute_diluted_shares_tsm() — basic shares + in-the-money options (TSM) + warrants + PSUs')
ref('Parts 3.6, 44')
item('Implement rollforward_basic_shares() — opening + new_issuances - buybacks = closing')
ref('Part 3.6')

h3('2I — DuPont & Profitability Analytics')
item('Implement compute_dupont_3factor() — ROE = Net_Margin × Asset_Turnover × Equity_Multiplier')
ref('Part 22.1')
item('Implement compute_dupont_5factor() — ROE = Tax_Burden × Interest_Burden × EBIT_Margin × Asset_Turnover × Leverage')
ref('Part 22.2')
item('Implement compute_roic() — NOPAT / Invested_Capital (average)')
ref('Parts 32, A.6, 52')
item('Implement compute_incremental_roic() — delta_NOPAT / delta_InvestedCapital')
ref('Part 69')
item('Implement compute_eva() and compute_eva_series() — EVA = NOPAT - (WACC × invested_capital)')
ref('Part 74')
item('Implement compute_fcfe() — UFCF - Interest×(1-t) + net_new_debt')
ref('Part 33.2')
item('Implement compute_cfads() — UFCF - mandatory_debt_service (for lender coverage ratios)')
ref('Part 18')
item('Implement coverage_ratios() — DSCR, ICR, FCCR, leverage ratio suite')
ref('Part 18.2')
item('Implement compute_bvps() — total equity / basic shares; P/B ratio')
ref('Part 47.2')

# ═══════════════════════════════════════════════════════════════════════════
h2('PHASE 3 — WACC & Discount Rate Engine')
ref('Parts: 4.3, 4.4, 38, 46.3, 48.1, 50, 71, 72, 73, 75')
# ═══════════════════════════════════════════════════════════════════════════

item('Implement fetch_industry_beta() — load from Damodaran industry beta table by GICS sector')
ref('Part 4.3, A.4')
item('Implement compute_unlevered_beta() — unlever peer betas: Bu = Bl / (1 + (1-t) × D/E)')
ref('Part 4.3')
item('Implement compute_relevered_beta() — re-lever to subject D/E: Bl = Bu × (1 + (1-t) × D/E)')
ref('Part 4.3')
item('Implement compute_unlevered_beta_cash_adjusted() — Damodaran: B_op = B_lev / (1 - cash/firm_value)')
ref('Part 71')
item('Implement compute_predicted_beta_blume() — Blume: B_adj = 0.67 × B_raw + 0.33 × 1.0 (mean reversion)')
ref('Source: Macabacus WACC page; standard Blume (1975) adjustment')
item('Implement compute_total_beta() — Total_Beta = Market_Beta / correlation (for private company ke)')
ref('Part 73')
item('Implement compute_cost_of_equity() — CAPM: ke = rf + beta_relevered × ERP + size_premium + CRP')
ref('Parts 4.3, 38')
item('Implement fetch_size_premium() — Duff & Phelps / Kroll CRSP decile lookup by market cap')
ref('Part 38.1')
item('Implement fetch_crp() — Damodaran CRP table lookup by country (for non-US companies)')
ref('Part 38.2')
item('Implement compute_cost_of_debt() — yield_to_maturity from FMP, or LIBOR/SOFR + spread; apply after-tax')
ref('Part 50.1')
item('Implement compute_wacc() — ke×(E/V) + kd_at×(D/V) using market-value weights at target capital structure')
ref('Parts 4.3, 4.4, 50.1')
item('Implement compute_wacc_with_preferred() — 4-component: ke×E + kd_at×D + rp×P + k_lease×L (all / V)')
ref('Parts 75, source: Macabacus WACC')
item('Implement compute_wacc_with_leases() — IFRS 16: add lease component kl_at = lease_rate×(1-t)×(L/V)')
ref('Part 75')
item('Implement compute_cross_currency_wacc() — ensure Rf and ERP are in same currency; add CRP for non-US')
ref('Part 46.3')
item('Implement wacc_mean_reversion_schedule() — fade WACC from current toward long-run average over forecast period')
ref('Part 48.1')

# ═══════════════════════════════════════════════════════════════════════════
h2('PHASE 4 — DCF Engine: Discounting, Terminal Value & EV Bridge')
ref('Parts: 3.2, 3.3, 3.4, 17, 41, 52, 65, 72')
# ═══════════════════════════════════════════════════════════════════════════

item('Implement compute_pv_ufcfs() — PV each forecast UFCF using mid-year convention: PV = UFCF / (1+WACC)^(t-0.5)')
ref('Parts 3.2, 4.5')
item('Implement compute_terminal_value_gordon() — TV = FCFn × (1+g) / (WACC - g); validate g < nominal GDP growth')
ref('Parts 3.3, 41.1, C.1')
item('Implement compute_terminal_value_exit_multiple() — TV = LTM_EBITDA_terminal × exit_multiple')
ref('Part 3.3')
item('Implement compute_reinvestment_rate() — RR = growth / ROIC; validate ROIC > WACC in terminal year')
ref('Parts 52.1, 52.2')
item('Implement compute_tv_nopat_reinvestment() — TV = NOPAT_terminal × (1 - RR) / (WACC - g) (simplified TV)')
ref('Part 52.2')
item('Implement validate_terminal_roic() — assert terminal ROIC > terminal WACC; warn if close to zero NPV')
ref('Part 32.1')
item('Implement validate_reinvestment_consistency() — assert implied g from reinvestment rate matches terminal g')
ref('Part 69.2')
item('Implement compute_pv_terminal_value() — discount TV to present using end-of-period convention')
ref('Part 3.2')
item('Implement compute_enterprise_value() — sum PV(UFCFs) + PV(TV)')
ref('Part 3')
item('Implement compute_xnpv() — exact-date DCF using actual day counts (XNPV convention)')
ref('Part 72')
item('Implement compute_xirr() — implied IRR from exact-date cash flows using brentq solver')
ref('Part 72')
item('Implement compute_apv() — APV = Unlevered_NPV + PV_ITS + PV_TLC; ITS_used = min(ITS, taxes_paid)')
ref('Part 17.2')
item('Implement compute_equity_value_per_share() — EV bridge: EV - net_debt - preferred - NCI + equity_investments = equity value / diluted shares')
ref('Parts 3.4, 34')

h3('EV Bridge Components')
item('Deduct: total interest-bearing debt (short-term + current LT + long-term + finance leases)')
ref('Part 15.1')
item('Deduct: operating lease liabilities (IFRS 16 / ASC 842 balance sheet amount)')
item('Deduct: preferred stock at liquidation value')
item('Deduct: NCI / minority interest (book value of noncontrolling equity)')
item('Deduct: pension underfunding (PBO - fair value of plan assets, net of deferred tax)')
item('Deduct: contingent liabilities / undrawn commitments (if material and disclosed)')
item('Add: cash and cash equivalents')
item('Add: short-term investments')
item('Add: equity method investments (associates at book value if not in UFCF)')
item('Add: net operating loss carryforwards (PV of tax benefit, if material)')
ref('Part 34')

# ═══════════════════════════════════════════════════════════════════════════
h2('PHASE 5 — Comparable Companies & Precedent Transactions')
ref('Parts: 5, 21, 26, 37, 46.2, 54, 77, 78')
# ═══════════════════════════════════════════════════════════════════════════

h3('5A — Trading Comps')
item('Implement select_peer_group() — auto-screen GICS sector/industry from FMP profile; return list of ticker peers')
ref('Parts 5.1, 46.2')
item('Implement fetch_peer_financials() — fetch LTM and NTM metrics for all peers')
ref('Parts 5.2, 37')
item('Implement compute_peer_ev() — market_cap + IBD - cash - ST_investments + NCI + preferred (per Macabacus)')
ref('Part 5.2')
item('Implement apply_ebitdar_adjustment() — for lease-heavy industries (retail, airlines), report EV/EBITDAR')
ref('Part 26')
item('Implement compute_peer_multiples() — EV/Revenue, EV/EBITDA, EV/EBIT, EV/EBITDAR, P/E, P/B, PEG')
ref('Parts 5.2, 26, 47.2, 54.4')
item('Implement exclude_nm_multiples() — mark as NM and exclude when denominator < 0 or multiple > 3×IQR from median')
ref('Part 21.2')
item('Implement compute_comps_summary_stats() — 25th percentile, median, 75th percentile for each multiple')
ref('Part 5.3')
item('Implement apply_multiples_to_subject() — derive implied EV range using peer 25th/median/75th × subject metrics')
ref('Part 5.3')
item('Implement check_peer_proforma_events() — FMP 8-K lookback; flag peers with recent material events')
ref('Part 78.2')
item('Implement apply_manual_proforma_adjustments() — from overrides/{TICKER}_comps_pf.json')
ref('Part 78.2')

h3('5B — Precedent Transactions')
item('Implement load_precedent_transactions() — from overrides/{TICKER}_transactions.json (user-supplied deal data)')
ref('Part 77.3')
item('Implement compute_transaction_multiples() — EV/LTM_Revenue and EV/LTM_EBITDA with NM outlier exclusion')
ref('Part 77.4')
item('Implement compute_transaction_comps_result() — apply deal multiples to subject; control-premium fallback if no deals')
ref('Part 77.4')

# ═══════════════════════════════════════════════════════════════════════════
h2('PHASE 6 — Sensitivity Analysis, Scenarios & Risk')
ref('Parts: 16, 36, 40, 47, 48, 65')
# ═══════════════════════════════════════════════════════════════════════════

item('Implement build_sensitivity_grid() — 12×9 grid: WACC ±1% (12 steps) × terminal growth ±0.5% (9 steps)')
ref('Parts 16, 6.2')
item('Implement compute_irr_implied_wacc() — find WACC that sets NPV = 0 (market-implied discount rate)')
ref('Part 16.3')
item('Implement run_scenario_analysis() — bull / base / bear scenarios; vary revenue growth, margins, WACC independently')
ref('Part 36')
item('Implement build_tornado_chart_data() — single-variable sensitivity: vary each assumption ±10%, rank by impact')
ref('Part 47.1')
item('Implement run_monte_carlo_dcf() — 10,000 simulations over revenue growth, EBIT margin, WACC, terminal growth; return distribution of equity values')
ref('Part 65')

# ═══════════════════════════════════════════════════════════════════════════
h2('PHASE 7 — Sector-Specific Modules')
ref('Parts: 26, 35, 59')
# ═══════════════════════════════════════════════════════════════════════════

item('Implement financial_company_gate() — detect GICS sector 40 (Financials); halt UFCF-DCF and output a clear error message')
ref('Part 35')
item('Implement reit_ffo_affo_model() — FFO = Net Income + D&A - gains_on_sale; AFFO = FFO - maintenance_capex')
ref('Part 59.1')
item('Implement ebitdar_adjustment() — for GICS 25 (Retail) and GICS 20301010 (Airlines): normalise for operating lease costs')
ref('Part 26')
item('Add stub: mining_nav_unsupported() — return UnsupportedError with explanation for mining/resource companies')
ref('Part 59.2')

# ═══════════════════════════════════════════════════════════════════════════
h2('PHASE 8 — Validation & Quality Control')
ref('Parts: 7, 24, 32, 40, 41, 42, 43, 44, 45, 46, 55, 61, 69, 70, 76')
# ═══════════════════════════════════════════════════════════════════════════

item('Implement check_ufcf_sign() — UFCF may be negative for early-stage companies; do not auto-clip')
item('Implement check_wacc_range() — assert 4% < WACC < 25%; warn outside 6–15%')
item('Implement check_terminal_growth_ceiling() — assert terminal g < 10-year nominal GDP growth estimate (~4-5%)')
ref('Parts 3.3, 41.1')
item('Implement check_tv_pct_of_ev() — warn if TV > 80% of total EV (typically 60–75% is normal range)')
ref('Part 41.1')
item('Implement check_sbc_terminal_dilution() — terminal UFCF must subtract expected SBC to avoid overstating TV')
ref('Part 41.2')
item('Implement check_revenue_growth_vs_margins() — if YoY revenue growth > 20% while EBIT margin falling, flag')
item('Implement check_capex_vs_da() — flag if CapEx < 50% of D&A (possible under-investment) or > 5× D&A (anomaly)')
item('Implement check_net_debt_sign() — negative net debt (net cash) is valid; do not force to zero')
item('Implement check_balance_sheet_closes() — assets = liabilities + equity within $1m tolerance (every year)')
ref('Part 76')
item('Implement check_nci_materiality() — NCI > 5% of NOPAT triggers compute_nopat_nci_adjusted()')
ref('Part 67')
item('Implement check_pension_materiality() — pension > 3% of EBIT triggers pension forecast module')
ref('Part 68.4')
item('Implement validate_reinvestment_consistency() — implied growth from ROIC × RR must match terminal g ± 0.5%')
ref('Part 69.2')
item('Implement check_terminal_roic_vs_wacc() — warn if terminal ROIC < WACC (value-destroying steady state)')
ref('Part 32.1')
item('Implement check_restatement_detection() — flag if same fiscal year appears twice in FMP data')
ref('Part 46.1')
item('Implement check_price_freshness() — warn if market price data > 3 trading days old')
ref('Part 55.1')
item('Run all validation checks after every model generation; log to audit trail and write to Validation sheet in Excel')
ref('Parts 7, 24, 39.1')

# ═══════════════════════════════════════════════════════════════════════════
h2('PHASE 9 — Excel Output Layer')
ref('Parts: 6, 16, 27, 30, 36, 39, 47, 54, 62, 63')
# ═══════════════════════════════════════════════════════════════════════════

h3('9A — Workbook Structure')
item('Create workbook with sheets in order: Cover | Model | Assumptions | Comps | Scenarios | Validation | README')
ref('Part 6.1')
item('Enable iterative calculation setting in openpyxl (for circular reference in interest expense)')
ref('Part 6.4')
item('Apply column widths, row heights, freeze panes (freeze row 1 and column A on Model sheet)')
ref('Parts 30, 54.3')
item('Apply named ranges for WACC, terminal growth, projection year cells')
ref('Part 6.5')
item('Apply number format codes: accounting $#,##0.0 for values; 0.0% for percentages; 0.0x for multiples')
ref('Part 62')
item('Set print area (Ctrl+P should show clean A4/Letter output)')
ref('Part 54.3')

h3('9B — Model Sheet')
item('Implement write_model_sheet() — rows per Part 30 column layout; LTM + 5 forecast years')
ref('Part 30')
item('Row group: Revenue (total, by segment if available), YoY growth %')
item('Row group: EBIT, EBIT margin %, Depreciation, EBITDA, EBITDA margin %')
item('Row group: NOPAT, Change in NOWC, CapEx, D&A, UFCF')
item('Row group: Discount factors (WACC, mid-year), PV of UFCFs')
item('Row group: Terminal value (Gordon growth and exit multiple), PV of TV')
item('Row group: Enterprise Value, EV bridge items, Equity Value, Diluted Shares, Price per Share')
item('Row group: WACC build (Rf, ERP, Beta, ke, kd, WACC)')
item('Row group: Key ratios (ROIC, ROCE, EV/EBITDA implied, P/E implied)')
item('Row group: Historical UFCF (back-test against actual FCF)')
ref('Part 32.3')
item('Apply golden rule: every historical cell is a formula pointing to raw data tab; never hard-code')
ref('Part 6.3')

h3('9C — Assumptions Sheet')
item('Implement write_assumptions_sheet() — list all model drivers with current values and editable override cells')
ref('Part 6.6')
item('Add data validation dropdowns where appropriate (e.g. scenario selector, terminal growth range)')
item('Add source/rationale column for every assumption')

h3('9D — Sensitivity Grid')
item('Implement write_sensitivity_sheet() — 12×9 WACC × terminal growth grid; color-coded by implied price vs current price')
ref('Part 16')

h3('9E — Comps / Peer Multiples Sheet')
item('Implement write_comps_sheet() — peer table: company name, market cap, EV, revenue, EBITDA, EV/EBITDA, EV/Rev, P/E')
ref('Part 54.1')
item('Add summary statistics row (high, mean, median, low) and subject company row for comparison')
item('Add pro forma warning rows below table if check_peer_proforma_events() flagged any peers')
ref('Part 78.3')
item('Add precedent transactions table below trading comps (from Part 77 data)')
ref('Part 77.5')
item('Add PEG ratio column and P/B column')
ref('Parts 54.4, 47.2')

h3('9F — Scenarios Sheet')
item('Implement write_scenarios_sheet() — bull / base / bear side-by-side: key assumptions + implied equity value per share')
ref('Part 36.3')

h3('9G — Football Field Chart (README Sheet)')
item('Implement build_ff_data_with_52wk() — 4-band football field builder: DCF | Trading Comps | Transaction Comps | 52-Wk Range')
ref('Parts 27, 79.3')
item('Implement write_football_field_data() — openpyxl stacked bar chart (floating bar technique)')
ref('Part 27.2')
item('Implement football_field_matplotlib() — optional high-quality PNG if matplotlib available')
ref('Part 27.3')
item('Add interpretation note below chart (DCF = minority basis; transaction comps = control premium included)')
ref('Part 79.4')
item('Plot current price as vertical line / diamond marker across all bands')
ref('Part 27')

h3('9H — Cover Page')
item('Implement write_cover_page() — company name, ticker, analysis date, analyst name, current price, implied equity value range, disclaimer')
ref('Part 54.2')

h3('9I — Validation Sheet')
item('Implement write_validation_sheet() — list all check names, PASS/WARN/FAIL status, actual value, threshold')
ref('Parts 7, 39.1')

h3('9J — Tornado Chart')
item('Implement write_tornado_chart() — horizontal bar chart showing sensitivity of equity value to each individual assumption')
ref('Part 47.1')

# ═══════════════════════════════════════════════════════════════════════════
h2('PHASE 10 — Configuration, CLI & Infrastructure')
ref('Parts: 9, 11, 31, 33, 39, 40, 47.3, 63, 66')
# ═══════════════════════════════════════════════════════════════════════════

item('Wire config hierarchy: CLI args → ticker override JSON → sector defaults → global defaults')
ref('Part 66')
item('Implement build_output_path() — output/TICKER_YYYY-MM-DD_vN.xlsx with timestamp and version suffix')
ref('Part 63')
item('Implement load_api_keys() — via python-dotenv; raise clear error if FMP_API_KEY missing')
ref('Part 39.2')
item('Implement error_recovery() — on partial API failure, log the error, use cached/fallback data, continue')
ref('Part 39.3')
item('Implement batch_mode() — iterate over list of tickers from CSV/JSON file; generate one workbook per ticker')
ref('Part 47.3')
item('Implement ValuationResult dataclass — single return object from main() with all key outputs')
ref('Part 70.1')

# ═══════════════════════════════════════════════════════════════════════════
h2('PHASE 11 — Testing')
ref('Parts: 49.2, 23, 24')
# ═══════════════════════════════════════════════════════════════════════════

item('Write unit tests for compute_ufcf() — test with known NIKE model inputs; assert within 1% of manual result')
item('Write unit tests for compute_wacc() — test with known beta, Rf, ERP, D/E inputs')
item('Write unit tests for compute_equity_value_per_share() — test EV bridge arithmetic')
item('Write unit tests for sensitivity grid — assert 12×9 array shape and monotonic relationship')
item('Write unit tests for all 20+ validation checks — test both passing and failing cases')
item('Write integration test (test_integration.py) — run full model on AAPL, MSFT, AMZN; assert output file created, equity value > 0, validation sheet shows PASS')
ref('Part 49.2')
item('Write regression test — run model on NIKE (NKE); compare equity value to historical model output within ±15%')
ref('Part 12.2')

# ═══════════════════════════════════════════════════════════════════════════
h2('PHASE 12 — Deployment & Delivery')
ref('Parts: 49, 56, 70')
# ═══════════════════════════════════════════════════════════════════════════

item('Write Dockerfile — Python 3.11-slim base; COPY requirements.txt; RUN pip install; COPY src; CMD [python, main.py]')
ref('Part 49.1')
item('Write docker-compose.yml — service: dcf_valuation; env_file: .env; volumes: ./output:/app/output')
ref('Part 49.1')
item('Implement deliver_by_email() — optional SMTP delivery of completed Excel to recipient list')
ref('Part 70.2')
item('Implement export_xlsx_to_pdf() — optional LibreOffice/unoconv export for PDF distribution')
ref('Part 49.3')
item('Implement webhook_notify() — optional POST to Slack/Teams webhook on completion or error')
ref('Part 49.3')
item('Final: run full model on 5 tickers (AAPL, MSFT, NKE, JPM=gated, O=REIT), review all outputs end-to-end')

# ═══════════════════════════════════════════════════════════════════════════
h2('IMPLEMENTATION ORDER SUMMARY')
# ═══════════════════════════════════════════════════════════════════════════

body(
    'Recommended build sequence — each phase unblocks the next:\n\n'
    '  Phase 0 → Phase 1 → Phase 2 → Phase 3 → Phase 4 → Phase 9A/B '
    '(minimal output) → Phase 5 → Phase 9C-J → Phase 6 → Phase 7 → '
    'Phase 8 → Phase 10 → Phase 11 → Phase 12\n\n'
    'MINIMUM VIABLE MODEL (can generate a working DCF): '
    'Phase 0 + Phase 1 (1A-1F) + Phase 2 (2A-2F) + Phase 3 + Phase 4 + Phase 9A-B.\n\n'
    'FULL PROFESSIONAL OUTPUT (IB-grade football field, comps, scenarios, validation): '
    'All 12 phases complete.'
)

body(
    'REFERENCE MAP — Part number to Phase:\n'
    '  Parts 1-12      → Phase 0-1 (architecture, data)\n'
    '  Parts 13-24     → Phase 2-3 (model mechanics, WACC)\n'
    '  Parts 25-34     → Phase 4-5 (DCF, EV bridge, comps)\n'
    '  Parts 35-49     → Phase 6-9 (scenarios, sectors, Excel)\n'
    '  Parts 50-65     → Phase 2-8 (advanced model components)\n'
    '  Parts 66-76     → Phase 2-8 (config, advanced adjustments)\n'
    '  Parts 77-79     → Phase 5-9 (precedent transactions, 52-wk, pro forma)'
)

# ─────────────────────────────────────────────────────────────────────────────
doc.save(DOCX_PATH)
import os
size = os.path.getsize(DOCX_PATH)
print(f'Saved. File size: {size:,} bytes ({size/1024:.1f} KB)')
print('Part 80 — Master Implementation Checklist appended.')
